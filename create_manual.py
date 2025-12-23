from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def create_installation_manual():
    """인프라 관계자용 설치 매뉴얼 생성"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('InstallTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('InstallH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('InstallH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('InstallH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    # 표지
    title = doc.add_paragraph('MSYS 설치 매뉴얼', style='InstallTitle')
    subtitle = doc.add_paragraph('인프라 구축 및 시스템 설치 가이드', style='InstallTitle')
    version = doc.add_paragraph('버전 1.14.2', style='InstallTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='InstallH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 시스템 요구사항 ........................ 3\n')
    toc.add_run('2. 네트워크 구성 ........................ 4\n')
    toc.add_run('3. 서버 환경 구축 ....................... 5\n')
    toc.add_run('4. 데이터베이스 설치 .................... 8\n')
    toc.add_run('5. 애플리케이션 배포 .................... 11\n')
    toc.add_run('6. 보안 설정 ........................... 14\n')
    toc.add_run('7. 모니터링 설정 ........................ 16\n')
    toc.add_run('부록 A. 설치 검증 체크리스트 ............. 18\n')

    doc.add_page_break()

    # 1. 시스템 요구사항
    doc.add_paragraph('1. 시스템 요구사항', style='InstallH1')
    reqs = doc.add_paragraph()
    reqs.add_run('1.1 하드웨어 요구사항\n').bold = True
    reqs.add_run('• CPU: Intel Xeon 또는 AMD EPYC, 2.0GHz 이상\n')
    reqs.add_run('• 메모리: 8GB 이상 (16GB 권장)\n')
    reqs.add_run('• 저장소: SSD 50GB 이상\n')
    reqs.add_run('• 네트워크: 1Gbps 이더넷\n\n')

    reqs.add_run('1.2 소프트웨어 요구사항\n').bold = True
    reqs.add_run('• 운영체제: Ubuntu Server 20.04 LTS 이상\n')
    reqs.add_run('• Python: 3.8 이상\n')
    reqs.add_run('• 데이터베이스: PostgreSQL 13 이상\n')
    reqs.add_run('• 웹 서버: Nginx (선택사항)\n\n')

    # 2. 네트워크 구성
    doc.add_paragraph('2. 네트워크 구성', style='InstallH1')
    network = doc.add_paragraph()
    network.add_run('2.1 방화벽 설정\n').bold = True
    network.add_run('필요한 포트 개방:\n')
    network.add_run('• 80 (HTTP)\n')
    network.add_run('• 443 (HTTPS)\n')
    network.add_run('• 5432 (PostgreSQL, 내부 통신만)\n')
    network.add_run('• 22 (SSH)\n\n')

    network.add_run('2.2 DNS 설정\n').bold = True
    network.add_run('도메인 및 SSL 인증서 구성\n\n')

    # 3. 서버 환경 구축
    doc.add_paragraph('3. 서버 환경 구축', style='InstallH1')

    # 3.1 Ubuntu Server 설치
    doc.add_paragraph('3.1 Ubuntu Server 설치', style='InstallH2')
    ubuntu = doc.add_paragraph()
    ubuntu.add_run('1. Ubuntu Server ISO 다운로드\n')
    ubuntu.add_run('2. 부팅 USB 생성\n')
    ubuntu.add_run('3. 서버에 설치 (기본 설정 사용)\n')
    ubuntu.add_run('4. 패키지 업데이트:\n\n')
    ubuntu.add_run('sudo apt update && sudo apt upgrade -y\n\n').font.name = 'Courier New'

    # 3.2 필수 패키지 설치
    doc.add_paragraph('3.2 필수 패키지 설치', style='InstallH2')
    packages = doc.add_paragraph()
    packages.add_run('sudo apt install -y python3 python3-pip python3-venv postgresql postgresql-contrib nginx curl wget\n\n').font.name = 'Courier New'

    # 4. 데이터베이스 설치
    doc.add_paragraph('4. 데이터베이스 설치', style='InstallH1')

    # 4.1 PostgreSQL 설정
    doc.add_paragraph('4.1 PostgreSQL 설정', style='InstallH2')
    pg = doc.add_paragraph()
    pg.add_run('1. PostgreSQL 서비스 시작:\n\n')
    pg.add_run('sudo systemctl start postgresql\n')
    pg.add_run('sudo systemctl enable postgresql\n\n').font.name = 'Courier New'

    pg.add_run('2. 데이터베이스 사용자 생성:\n\n')
    pg.add_run('sudo -u postgres psql\n').font.name = 'Courier New'
    pg.add_run('CREATE USER msys_user WITH PASSWORD \'strong_password\';\n').font.name = 'Courier New'
    pg.add_run('CREATE DATABASE msys_db OWNER msys_user;\n').font.name = 'Courier New'
    pg.add_run('\\q\n\n').font.name = 'Courier New'

    # 5. 애플리케이션 배포
    doc.add_paragraph('5. 애플리케이션 배포', style='InstallH1')

    # 5.1 애플리케이션 사용자 생성
    doc.add_paragraph('5.1 애플리케이션 사용자 생성', style='InstallH2')
    user = doc.add_paragraph()
    user.add_run('sudo useradd -m -s /bin/bash msys\n')
    user.add_run('sudo usermod -aG sudo msys\n\n').font.name = 'Courier New'

    # 5.2 소스 코드 배포
    doc.add_paragraph('5.2 소스 코드 배포', style='InstallH2')
    deploy = doc.add_paragraph()
    deploy.add_run('sudo -u msys mkdir -p /home/msys/app\n')
    deploy.add_run('sudo -u msys git clone <repository> /home/msys/app\n\n').font.name = 'Courier New'

    # 5.3 Python 환경 설정
    doc.add_paragraph('5.3 Python 환경 설정', style='InstallH2')
    python_env = doc.add_paragraph()
    python_env.add_run('cd /home/msys/app\n')
    python_env.add_run('python3 -m venv venv\n')
    python_env.add_run('source venv/bin/activate\n')
    python_env.add_run('pip install -r requirements.txt\n\n').font.name = 'Courier New'

    # 6. 보안 설정
    doc.add_paragraph('6. 보안 설정', style='InstallH1')
    security = doc.add_paragraph()
    security.add_run('6.1 방화벽 설정\n').bold = True
    security.add_run('sudo ufw allow 80\n')
    security.add_run('sudo ufw allow 443\n')
    security.add_run('sudo ufw --force enable\n\n').font.name = 'Courier New'

    security.add_run('6.2 SSL/TLS 설정\n').bold = True
    security.add_run('Let\'s Encrypt를 사용한 무료 SSL 인증서 설정\n\n')

    # 7. 모니터링 설정
    doc.add_paragraph('7. 모니터링 설정', style='InstallH1')
    monitoring = doc.add_paragraph()
    monitoring.add_run('7.1 로그 로테이션\n').bold = True
    monitoring.add_run('logrotate 설정으로 로그 파일 관리\n\n')

    monitoring.add_run('7.2 시스템 모니터링\n').bold = True
    monitoring.add_run('htop, iotop 등 모니터링 도구 설치\n\n')

    doc.save('MSYS_Installation_Manual.docx')
    print('설치 매뉴얼 생성 완료: MSYS_Installation_Manual.docx')

def create_function_manual():
    """개발자용 기능 매뉴얼 생성 - DB→DAO→SERVICE→JS→HTML 구조"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('FuncTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('FuncH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('FuncH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('FuncH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    heading4_style = doc.styles.add_style('FuncH4', WD_STYLE_TYPE.PARAGRAPH)
    heading4_style.font.size = Pt(11)
    heading4_style.font.bold = True

    code_style = doc.styles.add_style('FuncCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)

    # 표지
    title = doc.add_paragraph('MSYS 기능 매뉴얼', style='FuncTitle')
    subtitle = doc.add_paragraph('개발자용 데이터 흐름 참조 가이드', style='FuncTitle')
    version = doc.add_paragraph('버전 1.14.2', style='FuncTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='FuncH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 시스템 아키텍처 ........................ 3\n')
    toc.add_run('2. 인터페이스별 데이터 흐름 ................ 5\n')
    toc.add_run('   2.1 대시보드 인터페이스 .................. 5\n')
    toc.add_run('   2.2 데이터 분석 인터페이스 ................ 15\n')
    toc.add_run('   2.3 차트 분석 인터페이스 .................. 25\n')
    toc.add_run('   2.4 잔디 현황 인터페이스 .................. 35\n')
    toc.add_run('   2.5 매핑 관리 인터페이스 .................. 45\n')
    toc.add_run('   2.6 데이터 명세서 인터페이스 ................ 55\n')
    toc.add_run('3. 공통 컴포넌트 .......................... 65\n')
    toc.add_run('4. 설정 및 배포 .......................... 70\n')
    toc.add_run('부록 A. 데이터 흐름 다이어그램 ................ 75\n')

    doc.add_page_break()

    # 1. 시스템 아키텍처
    doc.add_paragraph('1. 시스템 아키텍처', style='FuncH1')
    arch = doc.add_paragraph()
    arch.add_run('1.1 데이터 흐름 개요\n').bold = True
    arch.add_run('MSYS는 다음과 같은 데이터 흐름을 따르는 계층형 아키텍처를 사용합니다:\n\n')
    arch.add_run('HTML → JavaScript → Service → DAO → Database\n\n')

    arch.add_run('1.2 각 계층의 역할\n').bold = True
    arch.add_run('• HTML: 사용자 인터페이스 및 데이터 표시\n')
    arch.add_run('• JavaScript: 사용자 상호작용 및 API 호출\n')
    arch.add_run('• Service: 비즈니스 로직 및 데이터 가공\n')
    arch.add_run('• DAO: 데이터베이스 쿼리 실행\n')
    arch.add_run('• Database: 데이터 저장 및 검색\n\n')

    # 2. 인터페이스별 데이터 흐름
    doc.add_paragraph('2. 인터페이스별 데이터 흐름', style='FuncH1')

    # 2.1 대시보드 인터페이스
    doc.add_paragraph('2.1 대시보드 인터페이스', style='FuncH2')
    dash_intro = doc.add_paragraph()
    dash_intro.add_run('대시보드는 시스템의 전체 현황을 실시간으로 모니터링하는 메인 인터페이스입니다.\n\n')

    # 대시보드 인터페이스 표
    dash_table = doc.add_table(rows=1, cols=5)
    dash_table.style = 'Table Grid'
    hdr_cells = dash_table.rows[0].cells
    hdr_cells[0].text = '레벨'
    hdr_cells[1].text = '파일/클래스'
    hdr_cells[2].text = '주요 메소드/함수'
    hdr_cells[3].text = '복잡도'
    hdr_cells[4].text = '설명'

    # DB 레벨 행
    db_row = dash_table.add_row().cells
    db_row[0].text = 'DB'
    db_row[1].text = 'DDL/*.sql\nsql/dashboard/*'
    db_row[2].text = '-'
    db_row[3].text = '고: 트리거\n서브쿼리 포함'
    db_row[4].text = '복합키, 타임존\nautomatic 로그'

    # DAO 레벨 메소드별 행들 (첫 번째 행에 파일 정보 표시)
    dao_methods = [
        ('get_summary(start_date, end_date, all_data, job_ids)', '중', '대시보드 요약 데이터 조회 - 지정된 기간과 Job에 대한 수집 통계 계산'),
        ('get_raw_data(start_date, end_date, job_ids, all_data)', '중', '원본 수집 데이터 조회 - 필터링된 수집 이력 데이터 반환'),
        ('get_analytics_success_rate_trend(start_date, end_date, job_ids, all_data)', '고', '성공률 추이 데이터 조회 - 시간별 성공률 변화 추이 계산'),
        ('get_trouble_by_code(start_date, end_date, job_ids, all_data)', '중', '장애 코드별 문제 데이터 조회 - 오류 유형별 발생 현황 분석'),
        ('get_event_log(start_date, end_date, all_data, allowed_job_ids)', '중', '이벤트 로그 조회 - 시스템 이벤트 및 변경 이력 확인'),
        ('get_daily_job_counts(job_id, start_date, end_date, all_data, job_ids)', '중', '일별 Job 카운트 조회 - Job별 일간 실행 통계 집계'),
        ('save_event(con_id, job_id, status, rqs_info)', '낮음', '이벤트 로그 저장 - 수집 이벤트 정보 기록'),
        ('get_distinct_job_ids(job_ids)', '낮음', '고유 Job ID 목록 조회 - 중복 제거된 Job ID 리스트 반환')
    ]

    # 첫 번째 메소드 행 (파일 정보 포함)
    first_dao_row = dash_table.add_row().cells
    first_dao_row[0].text = 'DAO'
    first_dao_row[1].text = 'mapper/dashboard_mapper.py\nsql/dashboard/*'
    first_dao_row[2].text = dao_methods[0][0]
    first_dao_row[3].text = dao_methods[0][1]
    first_dao_row[4].text = dao_methods[0][2]

    # 나머지 메소드 행들 (파일 정보 생략)
    for method, complexity, description in dao_methods[1:]:
        dao_row = dash_table.add_row().cells
        dao_row[0].text = ''  # 레벨 생략
        dao_row[1].text = ''  # 파일 정보 생략
        dao_row[2].text = method
        dao_row[3].text = complexity
        dao_row[4].text = description

    # SERVICE 레벨 메소드별 행들 (첫 번째 행에 파일 정보 표시)
    svc_methods = [
        ('get_summary(start_date, end_date, user)', '고', '대시보드 데이터 처리 메인 메소드 - 6단계 비즈니스 로직 실행'),
        ('get_raw_data(start_date, end_date, job_ids, all_data)', '중', '원본 수집 데이터 조회 및 기본 가공'),
        ('get_analytics_success_rate_trend(start_date, end_date, job_ids, all_data)', '고', '성공률 추이 데이터 계산 및 반환'),
        ('get_trouble_by_code(start_date, end_date, job_ids, all_data)', '중', '장애 코드별 문제 현황 분석'),
        ('get_event_log(start_date, end_date, all_data, allowed_job_ids)', '중', '시스템 이벤트 로그 조회'),
        ('get_daily_job_counts(job_id, start_date, end_date, all_data, job_ids)', '중', '일별 Job 실행 통계 계산'),
        ('_calculate_fail_streak(job_id)', '중', '특정 Job의 연속 실패 횟수 계산'),
        ('_get_allowed_job_ids(user)', '중', '사용자 권한에 따른 허용 Job ID 필터링'),
        ('_fetch_manager_settings_with_icons()', '중', '관리자 설정과 아이콘 정보 조회'),
        ('_combine_historical_and_today_data()', '고', '과거 데이터와 오늘 데이터 결합'),
        ('_apply_settings_and_filters()', '중', '설정 적용 및 추가 필터링'),
        ('_log_final_data_counts(processed_data)', '낮음', '최종 처리 데이터 로깅')
    ]

    # 첫 번째 메소드 행 (파일 정보 포함)
    first_svc_row = dash_table.add_row().cells
    first_svc_row[0].text = 'SERVICE'
    first_svc_row[1].text = 'service/dashboard_service.py'
    first_svc_row[2].text = svc_methods[0][0]
    first_svc_row[3].text = svc_methods[0][1]
    first_svc_row[4].text = svc_methods[0][2]

    # 나머지 메소드 행들 (파일 정보 생략)
    for method, complexity, description in svc_methods[1:]:
        svc_row = dash_table.add_row().cells
        svc_row[0].text = ''  # 레벨 생략
        svc_row[1].text = ''  # 파일 정보 생략
        svc_row[2].text = method
        svc_row[3].text = complexity
        svc_row[4].text = description

    # JS 레벨 행
    js_row = dash_table.add_row().cells
    js_row[0].text = 'JS'
    js_row[1].text = 'static/js/modules/dashboard/*'
    js_row[2].text = 'initializeDashboardData\nloadDashboardData\nrenderSummaryCards\nrenderStatusTable\nhandleDateFilterChange'
    js_row[3].text = '중: Promise.all\nAPI 호출\nDOM 조작'
    js_row[4].text = '병렬 로딩\n상태 추적\n이벤트 핸들링'

    # HTML 레벨 행
    html_row = dash_table.add_row().cells
    html_row[0].text = 'HTML'
    html_row[1].text = 'templates/dashboard.html'
    html_row[2].text = '-'
    html_row[3].text = '낮음'
    html_row[4].text = '템플릿 변수\nJinja2 필터'

    doc.add_paragraph('', style='FuncH3')  # 빈 줄

    # 실제 데이터 예시
    doc.add_paragraph('실제 데이터 예시:', style='FuncH3')
    example = doc.add_paragraph()
    example.add_run('• DB: job_id=\'CD101\', status=\'CD901\', start_dt=\'2025-12-18 14:30:15+09\'\n')
    example.add_run('• DAO: params=[\'2025-12-01\', \'2025-12-31\'], job_ids=[\'CD101\', \'CD102\']\n')
    example.add_run('• SERVICE: allowed_job_ids=[\'CD101\'], fail_streak=2, status=\'warning\'\n')
    example.add_run('• JS: dataFlowStatus.apiCallSuccess=true, loadedRecords=45\n')
    example.add_run('• HTML: {{ user.is_admin }} = true, {{ current_date }} = \'2025-12-18\'\n\n')

    # 2.2 데이터 분석 인터페이스
    doc.add_paragraph('2.2 데이터 분석 인터페이스', style='FuncH2')
    analysis_intro = doc.add_paragraph()
    analysis_intro.add_run('데이터 수집 결과를 심층 분석하고 필터링하여 조회하는 인터페이스입니다.\n\n')

    # 데이터 분석 인터페이스 표
    analysis_table = doc.add_table(rows=1, cols=5)
    analysis_table.style = 'Table Grid'
    hdr_cells = analysis_table.rows[0].cells
    hdr_cells[0].text = '레벨'
    hdr_cells[1].text = '파일/클래스'
    hdr_cells[2].text = '주요 메소드/함수'
    hdr_cells[3].text = '복잡도'
    hdr_cells[4].text = '설명'

    # DB 레벨 행
    analysis_db_row = analysis_table.add_row().cells
    analysis_db_row[0].text = 'DB'
    analysis_db_row[1].text = 'DDL/*.sql\nsql/analytics/*'
    analysis_db_row[2].text = '-'
    analysis_db_row[3].text = '고: 복합 조인\n다중 집계'
    analysis_db_row[4].text = '시간대 변환\n통계 계산'

    # DAO 레벨 메소드별 행들 (첫 번째 행에 파일 정보 표시)
    analysis_dao_methods = [
        ('get_user_access_stats(start_date, end_date, menu_id)', '중', '사용자별 메뉴 접근 통계 조회'),
        ('get_menu_access_stats(start_date, end_date, menu_id)', '중', '메뉴 접근 통계 조회'),
        ('get_menu_access_stats_weekly(start_date, end_date, menu_id)', '중', '주별 메뉴 접근 통계 조회'),
        ('get_yearly_total_stats(year, menu_id)', '중', '연간 총 통계 조회'),
        ('get_most_recent_data_date()', '낮음', '가장 최근 데이터 날짜 조회'),
        ('get_available_years_months()', '중', '사용 가능한 연월 목록 조회'),
        ('get_distinct_menu_names()', '낮음', '고유 메뉴 이름 목록 조회'),
        ('get_total_unique_users_by_week(start_date, end_date)', '중', '주별 고유 사용자 수 조회'),
        ('get_menu_access_stats_monthly(start_date, end_date, menu_id)', '중', '월별 메뉴 접근 통계 조회'),
        ('insert_user_access_log(user_id, menu_id, access_time)', '낮음', '사용자 접근 로그 기록'),
        ('get_menu_name_by_menu_id(menu_id)', '낮음', '메뉴 ID로 메뉴 이름 조회')
    ]

    # 첫 번째 메소드 행 (파일 정보 포함)
    first_analysis_dao_row = analysis_table.add_row().cells
    first_analysis_dao_row[0].text = 'DAO'
    first_analysis_dao_row[1].text = 'dao/analytics_dao.py'
    first_analysis_dao_row[2].text = analysis_dao_methods[0][0]
    first_analysis_dao_row[3].text = analysis_dao_methods[0][1]
    first_analysis_dao_row[4].text = analysis_dao_methods[0][2]

    # 나머지 메소드 행들 (파일 정보 생략)
    for method, complexity, description in analysis_dao_methods[1:]:
        analysis_dao_row = analysis_table.add_row().cells
        analysis_dao_row[0].text = ''  # 레벨 생략
        analysis_dao_row[1].text = ''  # 파일 정보 생략
        analysis_dao_row[2].text = method
        analysis_dao_row[3].text = complexity
        analysis_dao_row[4].text = description

    # SERVICE 레벨 메소드별 행들 (첫 번째 행에 파일 정보 표시)
    analysis_svc_methods = [
        ('get_analytics_data(start_date, end_date, filters)', '고', '분석 데이터 조회 및 가공'),
        ('process_chart_data(raw_data, chart_type)', '중', '차트 데이터 처리 및 변환'),
        ('calculate_statistics(data, stat_type)', '고', '통계 계산 수행'),
        ('generate_reports(data, report_format)', '고', '보고서 생성'),
        ('validate_date_range(start_date, end_date)', '낮음', '날짜 범위 유효성 검증'),
        ('apply_user_permissions(data, user)', '중', '사용자 권한 적용'),
        ('format_analytics_response(data)', '중', '분석 결과 포맷팅')
    ]

    # 첫 번째 메소드 행 (파일 정보 포함)
    first_analysis_svc_row = analysis_table.add_row().cells
    first_analysis_svc_row[0].text = 'SERVICE'
    first_analysis_svc_row[1].text = 'service/analysis_service.py'
    first_analysis_svc_row[2].text = analysis_svc_methods[0][0]
    first_analysis_svc_row[3].text = analysis_svc_methods[0][1]
    first_analysis_svc_row[4].text = analysis_svc_methods[0][2]

    # 나머지 메소드 행들 (파일 정보 생략)
    for method, complexity, description in analysis_svc_methods[1:]:
        analysis_svc_row = analysis_table.add_row().cells
        analysis_svc_row[0].text = ''  # 레벨 생략
        analysis_svc_row[1].text = ''  # 파일 정보 생략
        analysis_svc_row[2].text = method
        analysis_svc_row[3].text = complexity
        analysis_svc_row[4].text = description

    # JS 레벨 메소드별 행들 (첫 번째 행에 파일 정보 표시)
    analysis_js_methods = [
        ('initializeAnalysisPage()', '중', '데이터 분석 페이지 초기화'),
        ('fetchAndRenderAll()', '고', '전체 데이터 조회 및 렌더링'),
        ('fetchSummaryData(startDate, endDate)', '중', '요약 데이터 조회'),
        ('fetchTrendData(startDate, endDate)', '중', '추이 데이터 조회'),
        ('fetchRawData(startDate, endDate)', '중', '원본 데이터 조회'),
        ('renderSummaryCards(data)', '중', '요약 카드 렌더링'),
        ('renderTrendChart(data)', '중', '추이 차트 렌더링'),
        ('renderRawTable(data)', '중', '원본 데이터 테이블 렌더링'),
        ('renderJobInfoTable()', '중', 'Job 정보 테이블 렌더링'),
        ('initializeEventListeners()', '중', '이벤트 리스너 초기화'),
        ('askGeminiAnalysis()', '고', 'AI 분석 요청 및 결과 표시')
    ]

    # 첫 번째 메소드 행 (파일 정보 포함)
    first_analysis_js_row = analysis_table.add_row().cells
    first_analysis_js_row[0].text = 'JS'
    first_analysis_js_row[1].text = 'static/js/modules/data_analysis/*'
    first_analysis_js_row[2].text = analysis_js_methods[0][0]
    first_analysis_js_row[3].text = analysis_js_methods[0][1]
    first_analysis_js_row[4].text = analysis_js_methods[0][2]

    # 나머지 메소드 행들 (파일 정보 생략)
    for method, complexity, description in analysis_js_methods[1:]:
        analysis_js_row = analysis_table.add_row().cells
        analysis_js_row[0].text = ''  # 레벨 생략
        analysis_js_row[1].text = ''  # 파일 정보 생략
        analysis_js_row[2].text = method
        analysis_js_row[3].text = complexity
        analysis_js_row[4].text = description

    # HTML 레벨 행
    analysis_html_row = analysis_table.add_row().cells
    analysis_html_row[0].text = 'HTML'
    analysis_html_row[1].text = 'templates/data_analysis.html'
    analysis_html_row[2].text = '-'
    analysis_html_row[3].text = '중: 복합 템플릿'
    analysis_html_row[4].text = '필터 폼\n데이터 테이블\n차트 컨테이너\n탭 메뉴'

    doc.add_paragraph('', style='FuncH3')  # 빈 줄

    # 3. 공통 컴포넌트
    doc.add_paragraph('3. 공통 컴포넌트', style='FuncH1')
    common = doc.add_paragraph()
    common.add_run('3.1 데이터베이스 연결\n').bold = True
    common.add_run('파일: msys/database.py\n')
    common.add_run('함수: get_db_connection()\n\n')

    common.add_run('3.2 공통 DAO\n').bold = True
    common.add_run('파일: dao/__init__.py\n')
    common.add_run('기본 CRUD 메소드 제공\n\n')

    # 4. 설정 및 배포
    doc.add_paragraph('4. 설정 및 배포', style='FuncH1')
    deploy = doc.add_paragraph()
    deploy.add_run('4.1 환경 설정\n').bold = True
    deploy.add_run('.env 파일을 통한 환경별 설정 분리\n\n')

    deploy.add_run('4.2 배포 스크립트\n').bold = True
    deploy.add_run('Docker Compose를 활용한 컨테이너화 배포\n\n')

    doc.save('MSYS_Function_Manual_v8.docx')
    print('기능 매뉴얼 생성 완료: MSYS_Function_Manual_v8.docx')

def create_operation_manual():
    """운영자용 운영 매뉴얼 생성"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('OpTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('OpH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('OpH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('OpH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    # 표지
    title = doc.add_paragraph('MSYS 운영 매뉴얼', style='OpTitle')
    subtitle = doc.add_paragraph('시스템 운영 및 유지보수 가이드', style='OpTitle')
    version = doc.add_paragraph('버전 1.14.2', style='OpTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='OpH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 일상 운영 절차 ........................ 3\n')
    toc.add_run('2. 모니터링 ............................. 5\n')
    toc.add_run('3. 백업 및 복원 ........................ 8\n')
    toc.add_run('4. 장애 대응 ........................... 11\n')
    toc.add_run('5. 성능 최적화 ........................ 14\n')
    toc.add_run('6. 인증 시스템 ........................ 17\n')
    toc.add_run('7. 로그 관리 ........................... 20\n')
    toc.add_run('부록 A. 점검 체크리스트 .................. 23\n')

    doc.add_page_break()

    # 1. 일상 운영 절차
    doc.add_paragraph('1. 일상 운영 절차', style='OpH1')

    # 1.1 시스템 시작
    doc.add_paragraph('1.1 시스템 시작', style='OpH2')
    start_sys = doc.add_paragraph()
    start_sys.add_run('1.1.1 서버 접속\n').bold = True
    start_sys.add_run('SSH를 사용하여 서버에 접속합니다:\n\n')
    start_sys.add_run('ssh msys@server_ip_address\n\n').font.name = 'Courier New'

    start_sys.add_run('1.1.2 애플리케이션 디렉토리 이동\n').bold = True
    start_sys.add_run('MSYS 애플리케이션 디렉토리로 이동합니다:\n\n')
    start_sys.add_run('cd /home/msys/app\n\n').font.name = 'Courier New'

    start_sys.add_run('1.1.3 가상환경 활성화\n').bold = True
    start_sys.add_run('Python 가상환경을 활성화합니다:\n\n')
    start_sys.add_run('source venv/bin/activate\n\n').font.name = 'Courier New'

    start_sys.add_run('1.1.4 애플리케이션 실행\n').bold = True
    start_sys.add_run('MSYS 애플리케이션을 실행합니다:\n\n')
    start_sys.add_run('python msys_app.py\n\n').font.name = 'Courier New'
    start_sys.add_run('※ 애플리케이션이 백그라운드에서 실행되도록 하려면 nohup을 사용합니다:\n\n')
    start_sys.add_run('nohup python msys_app.py > /dev/null 2>&1 &\n\n').font.name = 'Courier New'

    # 1.2 시스템 중지
    doc.add_paragraph('1.2 시스템 중지', style='OpH2')
    stop_sys = doc.add_paragraph()
    stop_sys.add_run('1.2.1 정상적인 종료\n').bold = True
    stop_sys.add_run('실행 중인 프로세스를 찾아 종료합니다:\n\n')
    stop_sys.add_run('ps aux | grep python\n').font.name = 'Courier New'
    stop_sys.add_run('kill -TERM <process_id>\n\n').font.name = 'Courier New'

    stop_sys.add_run('1.2.2 강제 종료 (비상시)\n').bold = True
    stop_sys.add_run('응답이 없는 경우 강제 종료합니다:\n\n')
    stop_sys.add_run('kill -9 <process_id>\n\n').font.name = 'Courier New'

    # 1.3 서비스 상태 확인
    doc.add_paragraph('1.3 서비스 상태 확인', style='OpH2')
    check_svc = doc.add_paragraph()
    check_svc.add_run('1.3.1 프로세스 상태 확인\n').bold = True
    check_svc.add_run('Python 애플리케이션 프로세스 상태를 확인합니다:\n\n')
    check_svc.add_run('ps aux | grep python\n\n').font.name = 'Courier New'

    check_svc.add_run('1.3.2 데이터베이스 상태 확인\n').bold = True
    check_svc.add_run('PostgreSQL 서비스 상태를 확인합니다:\n\n')
    check_svc.add_run('systemctl status postgresql\n\n').font.name = 'Courier New'

    check_svc.add_run('1.3.3 웹 서비스 상태 확인\n').bold = True
    check_svc.add_run('웹 서버가 정상 응답하는지 확인합니다:\n\n')
    check_svc.add_run('curl -I http://localhost:5000\n\n').font.name = 'Courier New'

    # 2. 모니터링
    doc.add_paragraph('2. 모니터링', style='OpH1')

    # 2.1 대시보드 모니터링
    doc.add_paragraph('2.1 대시보드 모니터링', style='OpH2')
    dash_monitor = doc.add_paragraph()
    dash_monitor.add_run('웹 인터페이스를 통해 실시간으로 다음 항목을 모니터링합니다:\n')
    dash_monitor.add_run('• 데이터 수집 성공률\n')
    dash_monitor.add_run('• 시스템 상태\n')
    dash_monitor.add_run('• 오류 발생 현황\n\n')

    # 대시보드 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/dashboard_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.1: 대시보드 모니터링 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: dashboard_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    dash_desc = doc.add_paragraph()
    dash_desc.add_run('2.1.1 대시보드 화면 구성 요소 상세 설명\n').bold = True

    dash_desc.add_run('1. 요약 카드 영역:\n').bold = True
    dash_desc.add_run('• 전체 Job ID 개수: 시스템에 등록된 데이터 수집 작업의 총 개수\n')
    dash_desc.add_run('• 총 수집 건수: 선택된 기간 동안 성공적으로 수집된 데이터의 총 건수\n')
    dash_desc.add_run('• 평균 성공률: 모든 Job ID의 평균 데이터 수집 성공률 (백분율)\n\n')

    dash_desc.add_run('2. 상세 테이블:\n').bold = True
    dash_desc.add_run('• Job ID: 각 데이터 수집 작업의 고유 식별자 (예: "CD101", "CD102", "CD103", "CD104")\n')
    dash_desc.add_run('  - 시스템에서 자동으로 부여되는 작업 식별 번호\n')
    dash_desc.add_run('  - 각 작업을 구분하고 추적하는 데 사용\n')
    dash_desc.add_run('  - CD + 3자리 숫자 형식으로 구성\n\n')

    dash_desc.add_run('• 데이터명: 작업의 실제 데이터 수집 대상 이름\n')
    dash_desc.add_run('  - forecasts: 예보 데이터 수집\n')
    dash_desc.add_run('  - forecasts_o: 예보 데이터 수집 (외부)\n')
    dash_desc.add_run('  - live: 실시간 데이터 수집\n')
    dash_desc.add_run('  - live_o: 실시간 데이터 수집 (외부)\n\n')

    dash_desc.add_run('• 주기(cron): 데이터 수집 실행 주기\n')
    dash_desc.add_run('  - "0 6,10 * * *": 매일 6시, 10시에 실행\n')
    dash_desc.add_run('  - "0 16 */6 * *": 6일마다 16시에 실행\n')
    dash_desc.add_run('  - cron 표현식으로 자동화된 스케줄링\n\n')

    dash_desc.add_run('• 성공률: 해당 기간 동안의 데이터 수집 성공률 (백분율)\n')
    dash_desc.add_run('  - 기간별, Job ID별 통계를 모두 계산하여 산출\n')
    dash_desc.add_run('  - 성공한 수집 건수 ÷ 전체 수집 시도 건수 × 100\n')
    dash_desc.add_run('  - 오늘/주간/월간 등 다양한 기간 단위로 계산 가능\n\n')

    dash_desc.add_run('• 연속 실패 횟수: 최근 연속으로 실패한 데이터 수집 횟수\n')
    dash_desc.add_run('  - 마지막 성공 이후부터 현재까지 연속 실패한 시도 횟수\n')
    dash_desc.add_run('  - 0: 최근에 성공한 기록이 있음\n')
    dash_desc.add_run('  - 높을수록: 장기간 데이터 수집 실패 상태 지속\n\n')

    dash_desc.add_run('• 임계값: 설정된 성공률 기준값 (백분율)\n')
    dash_desc.add_run('  - 이 값 이상이면 정상, 미만이면 경고/위험 상태로 판단\n')
    dash_desc.add_run('  - 관리자가 각 Job별로 개별 설정 가능\n\n')

    dash_desc.add_run('• 상태: 현재 작업의 상태 표시 (정상/경고/위험)\n')
    dash_desc.add_run('  - 정상: 성공률 ≥ 임계값\n')
    dash_desc.add_run('  - 경고: 성공률 < 임계값이지만 심각하지 않은 상태\n')
    dash_desc.add_run('  - 위험: 성공률이 매우 낮거나 연속 실패가 많은 상태\n\n')

    dash_desc.add_run('3. 날짜 필터 컨트롤:\n').bold = True
    dash_desc.add_run('• 시작일/종료일: 모니터링할 기간을 선택\n')
    dash_desc.add_run('• 조회 버튼: 선택된 기간의 데이터를 다시 로드\n')
    dash_desc.add_run('• 기간별 보기: 일간/주간/월간 데이터 집계 방식 선택\n\n')

    dash_desc.add_run('4. 상태 표시 색상 코드 및 조건:\n').bold = True
    dash_desc.add_run('• 녹색 (●): 정상 상태 - (성공률 ≥ 임계값) AND (연속 실패 < 3)\n')
    dash_desc.add_run('• 노랑 (●): 경고 상태 - (성공률 80-95%) OR (연속 실패 3-5)\n')
    dash_desc.add_run('• 빨강 (●): 위험 상태 - (성공률 < 80%) OR (연속 실패 ≥ 5)\n\n')

    dash_desc.add_run('2.1.2 데이터 계산 방식\n').bold = True
    dash_desc.add_run('• 성공률 계산: (성공 건수 ÷ 총 시도 건수) × 100\n')
    dash_desc.add_run('• 연속 실패: 마지막 성공 이후 연속 실패 횟수\n')
    dash_desc.add_run('• 기간별 집계: 일별(당일)/주별(최근7일)/월별(최근30일)\n\n')

    dash_desc.add_run('2.1.3 실제 데이터 예시\n').bold = True
    dash_desc.add_run('• 정상 케이스: CD101 (성공률 98%, 연속 실패 0회) → 녹색 표시\n')
    dash_desc.add_run('• 경고 케이스: CD102 (성공률 85%, 연속 실패 2회) → 노랑 표시\n')
    dash_desc.add_run('• 위험 케이스: CD103 (성공률 65%, 연속 실패 7회) → 빨강 표시\n\n')

    dash_desc.add_run('※ 해당 값들은 관리자 설정에 의해 변경이 가능하며 현재 값들은 기본값을 기준으로 작성되었다.\n\n')

    # 2.2 데이터 분석 모니터링
    doc.add_paragraph('2.2 데이터 분석 모니터링', style='OpH2')
    data_analysis_monitor = doc.add_paragraph()
    data_analysis_monitor.add_run('데이터 수집 결과를 심층 분석하고 필터링을 통한 데이터 조회를 수행합니다.\n\n')

    # 데이터 분석 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/data_analysis_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.2: 데이터 분석 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: data_analysis_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    data_analysis_desc = doc.add_paragraph()
    data_analysis_desc.add_run('2.2.1 데이터 분석 화면 구성 요소 상세 설명\n').bold = True

    data_analysis_desc.add_run('1. 필터 컨트롤 영역:\n').bold = True
    data_analysis_desc.add_run('• 기간 선택: 분석할 데이터의 시작일과 종료일 설정\n')
    data_analysis_desc.add_run('• Job ID 선택: 분석할 특정 Job ID 선택 (다중 선택 가능)\n')
    data_analysis_desc.add_run('• 장애코드 필터: 특정 오류 코드로 데이터 필터링\n\n')

    data_analysis_desc.add_run('2. 데이터 조회 영역:\n').bold = True
    data_analysis_desc.add_run('• 요약 데이터: 선택된 기간의 전체 통계 요약\n')
    data_analysis_desc.add_run('• 추이 데이터: 시간에 따른 데이터 수집 추이 그래프\n')
    data_analysis_desc.add_run('• 원천 데이터: 실제 수집된 데이터의 상세 목록\n\n')

    data_analysis_desc.add_run('3. 데이터 분석 영역:\n').bold = True
    data_analysis_desc.add_run('• 데이터 필터링: Job ID, 장애코드, 기간별 필터링 기능\n')
    data_analysis_desc.add_run('• 분석 결과 표시: 필터링된 데이터의 통계 및 추이 분석\n\n')

    data_analysis_desc.add_run('2.2.2 데이터 구조 및 의미\n').bold = True
    data_analysis_desc.add_run('• 요약 데이터: 기간 내 총 수집 건수, 성공률, 실패 건수 등 집계 정보\n')
    data_analysis_desc.add_run('• 추이 데이터: 일별/시간별 데이터 수집 성공률 변화 추이\n')
    data_analysis_desc.add_run('• 원천 데이터: 각 수집 시도의 상세 정보 (시간, 상태, 오류 코드 등)\n\n')

    data_analysis_desc.add_run('2.2.3 사용자 인터랙션\n').bold = True
    data_analysis_desc.add_run('1. 필터 설정 후 [조회] 버튼 클릭하여 데이터 로드\n')
    data_analysis_desc.add_run('2. [AI 분석] 버튼 클릭하여 자동 분석 실행\n')
    data_analysis_desc.add_run('3. 분석 결과를 검토하고 필요한 조치 수행\n\n')

    data_analysis_desc.add_run('2.2.4 실제 데이터 예시\n').bold = True
    data_analysis_desc.add_run('• 정상 분석 케이스: CD101 (2025-12-01~2025-12-07) - 성공률 95%, 안정적인 수집 패턴\n')
    data_analysis_desc.add_run('• 문제 분석 케이스: CD102 (2025-12-01~2025-12-07) - 성공률 75%, 연속 실패 증가 추세\n\n')

    # 2.3 차트 분석 모니터링
    doc.add_paragraph('2.3 차트 분석 모니터링', style='OpH2')
    chart_analysis_monitor = doc.add_paragraph()
    chart_analysis_monitor.add_run('데이터 수집 성공률의 시간적 추이와 장애 코드별 현황을 시각적인 차트로 분석합니다.\n\n')

    # 차트 분석 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/chart_analysis_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.3: 차트 분석 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: chart_analysis_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    chart_analysis_desc = doc.add_paragraph()
    chart_analysis_desc.add_run('2.3.1 차트 분석 화면 구성 요소 상세 설명\n').bold = True

    chart_analysis_desc.add_run('1. 필터 컨트롤 영역:\n').bold = True
    chart_analysis_desc.add_run('• 기간 선택: 차트에 표시할 데이터의 시작일과 종료일 설정\n')
    chart_analysis_desc.add_run('• Job ID 선택: 분석할 Job ID 체크박스 (다중 선택 가능)\n')
    chart_analysis_desc.add_run('• 조회 버튼: 선택된 조건으로 차트 데이터 재로딩\n\n')

    chart_analysis_desc.add_run('2. 성공률 추이 차트:\n').bold = True
    chart_analysis_desc.add_run('• 시간에 따른 각 Job의 성공률 변화 그래프\n')
    chart_analysis_desc.add_run('• X축: 날짜/시간, Y축: 성공률 (백분율)\n')
    chart_analysis_desc.add_run('• 범례: 각 Job ID별 색상 구분\n\n')

    chart_analysis_desc.add_run('3. 장애 코드별 현황 차트:\n').bold = True
    chart_analysis_desc.add_run('• 선택된 기간 내 장애 코드 발생 빈도 파이 차트 또는 바 차트\n')
    chart_analysis_desc.add_run('• 각 장애 코드별 발생 건수 및 백분율 표시\n')
    chart_analysis_desc.add_run('• 상세 정보: 마우스 오버 시 구체적인 수치 표시\n\n')

    chart_analysis_desc.add_run('2.3.2 데이터 구조 및 의미\n').bold = True
    chart_analysis_desc.add_run('• 성공률 추이: 일별/시간별 데이터 수집 성공률의 시계열 데이터\n')
    chart_analysis_desc.add_run('• 장애 코드 현황: 기간 내 각 오류 코드의 발생 빈도 통계\n')
    chart_analysis_desc.add_run('• 차트 데이터: Chart.js 라이브러리를 사용한 인터랙티브 시각화\n\n')

    chart_analysis_desc.add_run('2.3.3 사용자 인터랙션\n').bold = True
    chart_analysis_desc.add_run('1. 기간과 Job ID 선택 후 [조회] 버튼 클릭\n')
    chart_analysis_desc.add_run('2. 차트 위에 마우스 오버하여 상세 데이터 확인\n')
    chart_analysis_desc.add_run('3. 범례 클릭으로 특정 Job의 표시/숨김 토글\n\n')

    chart_analysis_desc.add_run('2.3.4 실제 데이터 예시\n').bold = True
    chart_analysis_desc.add_run('• 성공률 추이: CD101의 12월 한 달간 성공률 95-98% 안정적 유지\n')
    chart_analysis_desc.add_run('• 장애 코드 현황: ERR001 (50%), ERR002 (30%), 기타 (20%)\n\n')

    # 2.4 잔디 현황 모니터링
    doc.add_paragraph('2.4 잔디 현황 모니터링', style='OpH2')
    jandi_monitor = doc.add_paragraph()
    jandi_monitor.add_run('Github의 활동 그래프처럼 각 Job의 일별 데이터 수집 현황을 히트맵 형태로 시각화하여 보여줍니다.\n\n')

    # 잔디 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/jandi_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.4: 잔디 현황 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: jandi_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    jandi_desc = doc.add_paragraph()
    jandi_desc.add_run('2.4.1 잔디 현황 화면 구성 요소 상세 설명\n').bold = True

    jandi_desc.add_run('1. 기간 선택 컨트롤:\n').bold = True
    jandi_desc.add_run('• 시작일/종료일: 히트맵에 표시할 기간 설정\n')
    jandi_desc.add_run('• 조회 버튼: 선택된 기간의 데이터 재로딩\n\n')

    jandi_desc.add_run('2. Job 목록 테이블:\n').bold = True
    jandi_desc.add_run('• Job ID: 각 데이터 수집 작업의 고유 식별자\n')
    jandi_desc.add_run('• Job 이름: 작업의 표시 이름\n')
    jandi_desc.add_run('• 펼치기/접기 버튼: 히트맵 표시 토글\n\n')

    jandi_desc.add_run('3. 히트맵 영역:\n').bold = True
    jandi_desc.add_run('• 일별 활동 표시: 각 날짜의 수집 건수에 따른 색상 표시\n')
    jandi_desc.add_run('• 색상 범례: 수집 건수별 색상 구분 (낮음 → 높음)\n')
    jandi_desc.add_run('• 마우스 오버: 구체적인 날짜와 수집 건수 표시\n\n')

    jandi_desc.add_run('2.4.2 데이터 구조 및 의미\n').bold = True
    jandi_desc.add_run('• 히트맵 데이터: 일별 수집 성공 건수 기반 색상 매핑\n')
    jandi_desc.add_run('• 색상 등급: 0건(흰색) → 1-5건(연한 녹색) → 6-10건(진한 녹색) → 10건+(어두운 녹색)\n')
    jandi_desc.add_run('• 데이터 집계: 선택된 기간 내 일별 총 수집 건수\n\n')

    jandi_desc.add_run('2.4.3 상태 표시 기준 및 색상 코드\n').bold = True
    jandi_desc.add_run('• 활동 없음: 흰색 (수집 건수 = 0)\n')
    jandi_desc.add_run('• 낮은 활동: 연한 녹색 (수집 건수 1-5)\n')
    jandi_desc.add_run('• 보통 활동: 중간 녹색 (수집 건수 6-10)\n')
    jandi_desc.add_run('• 높은 활동: 진한 녹색 (수집 건수 11+)\n\n')

    jandi_desc.add_run('2.4.4 실제 데이터 예시\n').bold = True
    jandi_desc.add_run('• 정상 패턴: CD101의 12월 히트맵 - 매일 8-12건 수집 (진한 녹색)\n')
    jandi_desc.add_run('• 불규칙 패턴: CD102의 12월 히트맵 - 주말 낮은 활동, 평일 높은 활동\n')
    jandi_desc.add_run('• 문제 패턴: CD103의 12월 히트맵 - 중간에 연속 흰색 구간 (수집 실패)\n\n')

    jandi_desc.add_run('2.4.5 사용자 인터랙션\n').bold = True
    jandi_desc.add_run('1. 기간 선택 후 [조회] 버튼 클릭하여 Job 목록 로드\n')
    jandi_desc.add_run('2. 관심 있는 Job의 [펼치기] 버튼 클릭하여 히트맵 표시\n')
    jandi_desc.add_run('3. 히트맵 위 마우스 오버하여 상세 데이터 확인\n')
    jandi_desc.add_run('4. 색상 범례로 활동 강도 파악\n\n')

    # 2.5 매핑 관리 모니터링
    doc.add_paragraph('2.5 매핑 관리 모니터링', style='OpH2')
    mapping_monitor = doc.add_paragraph()
    mapping_monitor.add_run('데이터베이스 테이블 간의 컬럼 매핑 정보를 관리하고 조회하는 기능을 제공합니다.\n\n')

    # 매핑 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/mapping_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.5: 매핑 관리 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: mapping_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    mapping_desc = doc.add_paragraph()
    mapping_desc.add_run('2.5.1 매핑 관리 화면 구성 요소 상세 설명\n').bold = True

    mapping_desc.add_run('1. 매핑 조회 영역:\n').bold = True
    mapping_desc.add_run('• 모든 매핑 보기: 시스템의 모든 매핑 정보 표시\n')
    mapping_desc.add_run('• 미매핑 컬럼 보기: 아직 매핑되지 않은 컬럼 목록\n')
    mapping_desc.add_run('• 새 매핑 추가: 새로운 매핑 관계 생성\n\n')

    mapping_desc.add_run('2. 매핑 테이블:\n').bold = True
    mapping_desc.add_run('• 소스 테이블: 원본 데이터 테이블 이름\n')
    mapping_desc.add_run('• 소스 컬럼: 원본 테이블의 컬럼명\n')
    mapping_desc.add_run('• 타겟 테이블: 대상 데이터 테이블 이름\n')
    mapping_desc.add_run('• 타겟 컬럼: 대상 테이블의 컬럼명\n')
    mapping_desc.add_run('• 매핑 유형: 매핑 관계의 종류 (직접, 변환, 계산 등)\n\n')

    mapping_desc.add_run('3. 매핑 작업 버튼:\n').bold = True
    mapping_desc.add_run('• 추가: 새로운 매핑 관계 생성\n')
    mapping_desc.add_run('• 수정: 기존 매핑 관계 편집\n')
    mapping_desc.add_run('• 삭제: 매핑 관계 제거\n')
    mapping_desc.add_run('• 내보내기: 매핑 정보를 파일로 저장\n\n')

    mapping_desc.add_run('2.5.2 데이터 구조 및 의미\n').bold = True
    mapping_desc.add_run('• 매핑 정보: 테이블 간 컬럼 관계 정의 데이터\n')
    mapping_desc.add_run('• 소스/타겟: 데이터 변환의 입력과 출력 관계\n')
    mapping_desc.add_run('• 매핑 유형: 데이터 변환 방식 (1:1, 1:N, 계산식 등)\n\n')

    mapping_desc.add_run('2.5.3 사용자 인터랙션\n').bold = True
    mapping_desc.add_run('1. [모든 매핑 보기] 또는 [미매핑 컬럼 보기] 선택\n')
    mapping_desc.add_run('2. 매핑할 컬럼 선택 후 [새 매핑 추가] 클릭\n')
    mapping_desc.add_run('3. 소스와 타겟 컬럼 관계 설정\n')
    mapping_desc.add_run('4. [저장] 버튼으로 매핑 관계 저장\n\n')

    mapping_desc.add_run('2.5.4 실제 데이터 예시\n').bold = True
    mapping_desc.add_run('• 직접 매핑: source_table.id → target_table.user_id\n')
    mapping_desc.add_run('• 변환 매핑: source_table.status_code → target_table.status (코드 변환)\n')
    mapping_desc.add_run('• 계산 매핑: source_table.price * 1.1 → target_table.final_price\n\n')

    # 2.6 데이터 명세서 모니터링
    doc.add_paragraph('2.6 데이터 명세서 모니터링', style='OpH2')
    data_spec_monitor = doc.add_paragraph()
    data_spec_monitor.add_run('시스템에서 사용하는 외부 데이터(API 등)의 명세를 관리하는 기능을 제공합니다.\n\n')

    # 데이터 명세서 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/data_spec_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.6: 데이터 명세서 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: data_spec_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    data_spec_desc = doc.add_paragraph()
    data_spec_desc.add_run('2.6.1 데이터 명세서 화면 구성 요소 상세 설명\n').bold = True

    data_spec_desc.add_run('1. 명세서 목록:\n').bold = True
    data_spec_desc.add_run('• 명세서 이름: 각 데이터 소스의 식별자\n')
    data_spec_desc.add_run('• 설명: 명세서에 대한 간단한 설명\n')
    data_spec_desc.add_run('• URL: 데이터 소스의 접근 주소\n')
    data_spec_desc.add_run('• 최종 수정일: 명세서 마지막 업데이트 일시\n\n')

    data_spec_desc.add_run('2. 명세서 작업 버튼:\n').bold = True
    data_spec_desc.add_run('• 새 명세서: 새로운 데이터 명세서 생성\n')
    data_spec_desc.add_run('• URL에서 가져오기: 웹 주소로부터 명세서 자동 생성\n')
    data_spec_desc.add_run('• 수정: 기존 명세서 편집\n')
    data_spec_desc.add_run('• 삭제: 명세서 제거\n\n')

    data_spec_desc.add_run('3. 명세서 상세 정보:\n').bold = True
    data_spec_desc.add_run('• API 엔드포인트: 데이터 접근 주소\n')
    data_spec_desc.add_run('• 요청 방식: HTTP 메소드 (GET, POST 등)\n')
    data_spec_desc.add_run('• 파라미터: API 호출에 필요한 매개변수\n')
    data_spec_desc.add_run('• 응답 형식: 반환 데이터 구조 (JSON, XML 등)\n\n')

    data_spec_desc.add_run('2.6.2 데이터 구조 및 의미\n').bold = True
    data_spec_desc.add_run('• 명세서 메타데이터: API 이름, 설명, 접근 권한 등\n')
    data_spec_desc.add_run('• API 스펙: 엔드포인트, 메소드, 파라미터, 응답 형식\n')
    data_spec_desc.add_run('• 데이터 스키마: 요청/응답 데이터 구조 정의\n\n')

    data_spec_desc.add_run('2.6.3 사용자 인터랙션\n').bold = True
    data_spec_desc.add_run('1. [새 명세서] 클릭하여 새로운 명세서 생성 시작\n')
    data_spec_desc.add_run('2. 또는 [URL에서 가져오기]로 자동 생성\n')
    data_spec_desc.add_run('3. 명세서 정보 입력 (이름, URL, 설명 등)\n')
    data_spec_desc.add_run('4. [저장]으로 명세서 등록 완료\n\n')

    data_spec_desc.add_run('2.6.4 실제 데이터 예시\n').bold = True
    data_spec_desc.add_run('• 날씨 API: 기상청 API 명세서 (GET /weather, 파라미터: location, date)\n')
    data_spec_desc.add_run('• 사용자 API: 외부 시스템 사용자 정보 (POST /users, 인증 토큰 필요)\n')
    data_spec_desc.add_run('• 금융 API: 주식 정보 조회 (GET /stocks/{symbol}, 실시간 데이터)\n\n')

    # 2.7 관리자 설정 모니터링
    doc.add_paragraph('2.7 관리자 설정 모니터링', style='OpH2')
    admin_monitor = doc.add_paragraph()
    admin_monitor.add_run('시스템의 각종 설정을 관리하고 모니터링하는 기능을 제공합니다.\n\n')

    # 관리자 설정 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/admin_basic_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.7: 관리자 설정 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_basic_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    admin_desc = doc.add_paragraph()
    admin_desc.add_run('2.7.1 관리자 설정 화면 구성 요소 상세 설명\n').bold = True

    admin_desc.add_run('1. 탭 메뉴:\n').bold = True
    admin_desc.add_run('화면 상단의 탭을 클릭하여 다른 설정 화면으로 이동할 수 있습니다.\n')
    admin_desc.add_run('• 기본 설정 (현재 선택됨)\n')
    admin_desc.add_run('• 수집 스케줄 설정\n')
    admin_desc.add_run('• Icon 관리\n')
    admin_desc.add_run('• 차트/시각화 설정\n')
    admin_desc.add_run('• 사용자 관리\n')
    admin_desc.add_run('• 데이터 접근 권한\n')
    admin_desc.add_run('• 통계\n\n')

    admin_desc.add_run('2. 설정 테이블:\n').bold = True
    admin_desc.add_run('Job ID별 설정을 표 형식으로 표시합니다.\n')
    admin_desc.add_run('• Job ID: 작업 고유 식별자\n')
    admin_desc.add_run('• Job 이름: 작업의 표시 이름\n')
    admin_desc.add_run('• Job 설명: 작업에 대한 설명\n')
    admin_desc.add_run('• 비고: 추가 메모\n\n')

    admin_desc.add_run('3. 연속 실패시 표시방법:\n').bold = True
    admin_desc.add_run('데이터 수집이 연속으로 실패할 경우의 표시 설정입니다.\n')
    admin_desc.add_run('• 실패(CNT): 연속 실패 횟수 임계값\n')
    admin_desc.add_run('• 실패(Icon): 실패 상태 아이콘 선택\n')
    admin_desc.add_run('• 실패(Color): 실패 상태 색상 선택\n')
    admin_desc.add_run('• 경고(CNT): 경고 상태 임계값\n')
    admin_desc.add_run('• 경고(Icon): 경고 상태 아이콘 선택\n')
    admin_desc.add_run('• 경고(Color): 경고 상태 색상 선택\n')
    admin_desc.add_run('• 정상(Icon): 정상 상태 아이콘 선택\n')
    admin_desc.add_run('• 정상(Color): 정상 상태 색상 선택\n\n')

    admin_desc.add_run('4. 성공률 기준 임계치 설정:\n').bold = True
    admin_desc.add_run('각 기간별 성공률 임계값을 설정합니다.\n')
    admin_desc.add_run('• 일간 임계값: 일일 성공률 기준\n')
    admin_desc.add_run('• 주간 임계값: 주간 성공률 기준\n')
    admin_desc.add_run('• 월간 임계값: 월간 성공률 기준\n')
    admin_desc.add_run('• 반기 임계값: 반기 성공률 기준\n')
    admin_desc.add_run('• 연간 임계값: 연간 성공률 기준\n')
    admin_desc.add_run('• 성공 아이콘/색상: 성공 상태 표시\n')
    admin_desc.add_run('• 경고 아이콘/색상: 경고 상태 표시\n\n')

    admin_desc.add_run('5. 대시보드 표시 여부:\n').bold = True
    admin_desc.add_run('해당 Job ID의 데이터를 대시보드에 표시할지 여부를 설정합니다.\n\n')

    admin_desc.add_run('6. 컬러 팔레트:\n').bold = True
    admin_desc.add_run('색상 선택 시 사용할 수 있는 미리 정의된 색상 팔레트입니다.\n\n')

    admin_desc.add_run('7. 작업 버튼:\n').bold = True
    admin_desc.add_run('• 행 추가: 새로운 Job ID 설정 행을 추가합니다.\n')
    admin_desc.add_run('• 기본 설정 저장: 현재 설정을 저장합니다.\n')
    admin_desc.add_run('• 설정 내보내기: 설정을 JSON 파일로 내보냅니다.\n')
    admin_desc.add_run('• 설정 가져오기: JSON 파일로부터 설정을 불러옵니다.\n\n')

    admin_desc.add_run('2.7.2 사용자 인터랙션\n').bold = True
    admin_desc.add_run('1. Job ID 선택: 설정할 Job ID의 행을 선택합니다.\n')
    admin_desc.add_run('2. 임계값 설정: 각 필드에 적절한 임계값을 입력합니다.\n')
    admin_desc.add_run('3. 설정 저장: [기본 설정 저장] 버튼을 클릭하여 변경사항을 저장합니다.\n')
    admin_desc.add_run('4. 설정 관리: 내보내기/가져오기 기능으로 설정을 백업/복원합니다.\n\n')

    admin_desc.add_run('2.7.3 실제 데이터 예시\n').bold = True
    admin_desc.add_run('• CD101 설정: 연속 실패 3회, 성공률 80%, 빨강/노랑/초록 색상\n')
    admin_desc.add_run('• CD102 설정: 연속 실패 2회, 성공률 75%, 사용자 정의 색상\n\n')

    # 2.8 시스템 리소스 모니터링
    doc.add_paragraph('2.8 시스템 리소스 모니터링', style='OpH2')
    sys_monitor = doc.add_paragraph()
    sys_monitor.add_run('시스템 리소스 사용량을 정기적으로 모니터링합니다:\n')
    sys_monitor.add_run('• CPU 사용률: 80% 미만 유지\n')
    sys_monitor.add_run('• 메모리 사용률: 85% 미만 유지\n')
    sys_monitor.add_run('• 디스크 사용률: 90% 미만 유지\n\n')

    sys_monitor.add_run('2.8.1 리소스 모니터링 명령어\n').bold = True
    sys_monitor.add_run('CPU 및 메모리 사용량 확인:\n\n')
    sys_monitor.add_run('top\n').font.name = 'Courier New'
    sys_monitor.add_run('htop\n\n').font.name = 'Courier New'

    sys_monitor.add_run('디스크 사용량 확인:\n\n')
    sys_monitor.add_run('df -h\n').font.name = 'Courier New'
    sys_monitor.add_run('du -sh /home/msys/app\n\n').font.name = 'Courier New'

    # 3. 백업 및 복원
    doc.add_paragraph('3. 백업 및 복원', style='OpH1')
    backup = doc.add_paragraph()
    backup.add_run('3.1 자동 백업 설정\n').bold = True
    backup.add_run('crontab에 다음 작업을 추가:\n\n')
    backup.add_run('0 2 * * * pg_dump -U msys_user msys_db > /backup/msys_$(date +\\%Y\\%m\\%d).sql\n\n').font.name = 'Courier New'

    backup.add_run('3.2 수동 백업\n').bold = True
    backup.add_run('pg_dump -U msys_user -h localhost msys_db > backup.sql\n\n').font.name = 'Courier New'

    backup.add_run('3.3 복원 절차\n').bold = True
    backup.add_run('1. 서비스 중지\n')
    backup.add_run('2. 데이터베이스 드롭 및 재생성\n')
    backup.add_run('3. 백업 파일로부터 복원:\n\n')
    backup.add_run('psql -U msys_user msys_db < backup.sql\n\n').font.name = 'Courier New'

    # 4. 장애 대응
    doc.add_paragraph('4. 장애 대응', style='OpH1')
    incident = doc.add_paragraph()
    incident.add_run('4.1 장애 유형별 대응\n\n').bold = True

    incident.add_run('웹 서비스 장애:\n').bold = True
    incident.add_run('1. 프로세스 상태 확인\n')
    incident.add_run('2. 로그 파일 검토\n')
    incident.add_run('3. 서비스 재시작\n\n')

    incident.add_run('데이터베이스 장애:\n').bold = True
    incident.add_run('1. PostgreSQL 서비스 상태 확인\n')
    incident.add_run('2. 연결 설정 검토\n')
    incident.add_run('3. 데이터베이스 복구\n\n')

    # 5. 성능 최적화
    doc.add_paragraph('5. 성능 최적화', style='OpH1')
    performance = doc.add_paragraph()
    performance.add_run('5.1 데이터베이스 최적화\n').bold = True
    performance.add_run('• 인덱스 생성 및 유지\n')
    performance.add_run('• 쿼리 최적화\n')
    performance.add_run('• 테이블 파티셔닝\n\n')

    performance.add_run('5.2 캐싱 전략\n').bold = True
    performance.add_run('• Redis 캐시 서버 도입 고려\n')
    performance.add_run('• 정적 파일 캐싱\n\n')

    # 6. 인증 시스템
    doc.add_paragraph('6. 인증 시스템', style='OpH1')

    # 6.1 로그인 화면
    doc.add_paragraph('6.1 로그인 화면', style='OpH2')
    login_desc = doc.add_paragraph()
    login_desc.add_run('시스템에 접근하기 위한 사용자 인증 화면입니다.\n\n')

    # 로그인 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/login_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 6.1: 로그인 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: login_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    login_elements = doc.add_paragraph()
    login_elements.add_run('6.1.1 화면 구성 요소\n').bold = True

    login_elements.add_run('1. 탭 메뉴:\n').bold = True
    login_elements.add_run('• 로그인: 사용자 인증을 위한 기본 탭\n')
    login_elements.add_run('• 회원가입: 신규 사용자 등록을 위한 탭\n\n')

    login_elements.add_run('2. 로그인 폼:\n').bold = True
    login_elements.add_run('• 사용자 ID: 시스템에 등록된 사용자 식별자 입력\n')
    login_elements.add_run('• 비밀번호: 사용자 인증을 위한 비밀번호 입력\n')
    login_elements.add_run('• 로그인 버튼: 인증 요청 실행\n')
    login_elements.add_run('• 비밀번호 초기화 요청: 비밀번호 분실 시 재설정 요청\n\n')

    login_elements.add_run('3. 시스템 안내:\n').bold = True
    login_elements.add_run('• 회원가입 후 관리자 승인이 필요하다는 안내 메시지\n')
    login_elements.add_run('• 문의사항 연락처 정보 표시\n\n')

    login_usage = doc.add_paragraph()
    login_usage.add_run('6.1.2 사용 방법\n').bold = True
    login_usage.add_run('1. 사용자 ID 입력: 시스템에 등록된 ID를 입력합니다.\n')
    login_usage.add_run('2. 비밀번호 입력: 해당 ID의 비밀번호를 입력합니다.\n')
    login_usage.add_run('3. 로그인 버튼 클릭: 인증을 요청합니다.\n')
    login_usage.add_run('4. 성공 시 대시보드로 자동 이동, 실패 시 오류 메시지 표시\n\n')

    # 6.2 회원가입 화면
    doc.add_paragraph('6.2 회원가입 화면', style='OpH2')
    register_desc = doc.add_paragraph()
    register_desc.add_run('신규 사용자가 시스템에 등록하기 위한 화면입니다.\n\n')

    # 회원가입 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/register_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 6.2: 회원가입 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: register_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    register_elements = doc.add_paragraph()
    register_elements.add_run('6.2.1 화면 구성 요소\n').bold = True

    register_elements.add_run('1. 회원가입 폼:\n').bold = True
    register_elements.add_run('• 사용자 ID: 영문/숫자 조합의 고유 식별자\n')
    register_elements.add_run('• 비밀번호: 보안 정책을 만족하는 비밀번호\n')
    register_elements.add_run('• 비밀번호 확인: 비밀번호 재입력으로 검증\n')
    register_elements.add_run('• 가입 신청 버튼: 회원가입 요청 실행\n\n')

    register_elements.add_run('2. 비밀번호 정책 체크리스트:\n').bold = True
    register_elements.add_run('실시간으로 비밀번호가 정책에 맞는지 검증하여 표시:\n')
    register_elements.add_run('• 8자 이상: 최소 길이 요구사항\n')
    register_elements.add_run('• 특수문자 1개 이상 포함: 보안 강도 향상\n')
    register_elements.add_run('• 연속된 숫자 사용 불가: 123, 456 등 패턴 방지\n')
    register_elements.add_run('• 동일 숫자 반복 불가: 111, 222 등 반복 방지\n\n')

    register_elements.add_run('3. 실시간 검증:\n').bold = True
    register_elements.add_run('• 입력 시 정책 준수 여부 실시간 표시 (✅/❌)\n')
    register_elements.add_run('• 비밀번호 확인 일치 여부 검증\n')
    register_elements.add_run('• 모든 조건 만족 시에만 가입 신청 버튼 활성화\n\n')

    register_usage = doc.add_paragraph()
    register_usage.add_run('6.2.2 사용 방법\n').bold = True
    register_usage.add_run('1. 사용자 ID 입력: 영문/숫자 조합으로 원하는 ID 입력\n')
    register_usage.add_run('2. 비밀번호 입력: 정책을 만족하는 비밀번호 입력\n')
    register_usage.add_run('3. 비밀번호 확인: 동일한 비밀번호 재입력\n')
    register_usage.add_run('4. 정책 체크리스트 확인: 모든 항목이 ✅ 표시되는지 확인\n')
    register_usage.add_run('5. 가입 신청 버튼 클릭: 회원가입 요청 (관리자 승인 대기)\n\n')

    # 6.4 회원가입 승인 및 권한 부여 흐름
    doc.add_paragraph('6.4 회원가입 승인 및 권한 부여 흐름', style='OpH2')
    approval_flow = doc.add_paragraph()
    approval_flow.add_run('신규 사용자의 시스템 접근 권한을 부여하는 전체 프로세스입니다.\n\n')

    approval_flow.add_run('6.4.1 권한 부여 단계별 흐름\n').bold = True
    approval_flow.add_run('1. 사용자 회원가입 신청\n').bold = True
    approval_flow.add_run('   - 사용자 ID, 비밀번호 입력\n')
    approval_flow.add_run('   - DB에 PENDING 상태로 저장\n')
    approval_flow.add_run('   - 관리자에게 승인 요청 알림\n\n')

    approval_flow.add_run('2. 관리자 승인 처리\n').bold = True
    approval_flow.add_run('   - 관리자 설정 → 사용자 관리 탭 접근\n')
    approval_flow.add_run('   - PENDING 상태 사용자 목록 확인\n')
    approval_flow.add_run('   - 승인 버튼 클릭 → APPROVED 상태로 변경\n')
    approval_flow.add_run('   - 비밀번호를 사용자 ID와 동일하게 초기화 (보안 정책 준수를 위해)\n')
    approval_flow.add_run('   - 사용자에게 초기 비밀번호 안내\n\n')

    approval_flow.add_run('2.1 사용자 첫 로그인 및 비밀번호 변경 강제\n').bold = True
    approval_flow.add_run('   - 사용자가 초기 비밀번호(사용자 ID와 동일)로 로그인 시도\n')
    approval_flow.add_run('   - 시스템이 보안 정책 위반 감지하여 비밀번호 변경 강제\n')
    approval_flow.add_run('   - "비밀번호를 변경해야 합니다. 초기화된 비밀번호는 안전하지 않습니다." 메시지 표시\n')
    approval_flow.add_run('   - 비밀번호 변경 화면으로 자동 리디렉션\n\n')

    approval_flow.add_run('3. 메뉴 권한 설정\n').bold = True
    approval_flow.add_run('   - 기본 권한 자동 부여: dashboard, collection_schedule\n')
    approval_flow.add_run('   - 추가 권한 수동 설정: data_analysis, chart_analysis 등\n')
    approval_flow.add_run('   - 권한별 메뉴 접근 제어\n\n')

    approval_flow.add_run('4. 데이터 접근 권한 설정\n').bold = True
    approval_flow.add_run('   - Job ID별 데이터 접근 권한 부여\n')
    approval_flow.add_run('   - 데이터 보안 및 프라이버시 제어\n')
    approval_flow.add_run('   - 권한에 따른 데이터 필터링 적용\n\n')

    approval_flow.add_run('6.4.2 권한 종류 및 의미\n').bold = True
    approval_flow.add_run('• dashboard: 대시보드 메뉴 접근 권한\n')
    approval_flow.add_run('• collection_schedule: 수집 일정 조회 권한\n')
    approval_flow.add_run('• data_analysis: 데이터 분석 메뉴 접근 권한\n')
    approval_flow.add_run('• chart_analysis: 차트 분석 메뉴 접근 권한\n')
    approval_flow.add_run('• jandi: 잔디 현황 메뉴 접근 권한\n')
    approval_flow.add_run('• mapping: 매핑 메뉴 접근 권한\n')
    approval_flow.add_run('• data_spec: 데이터 명세서 메뉴 접근 권한\n')
    approval_flow.add_run('• mngr_sett: 관리자 설정 메뉴 접근 권한\n\n')

    approval_flow.add_run('6.4.3 권한 설정 예시\n').bold = True
    approval_flow.add_run('• 일반 사용자: dashboard, collection_schedule, data_analysis\n')
    approval_flow.add_run('• 분석 담당자: 위 권한 + chart_analysis, jandi\n')
    approval_flow.add_run('• 관리자: 모든 권한 + mngr_sett\n\n')

    # 6.3 비밀번호 변경 화면
    doc.add_paragraph('6.3 비밀번호 변경 화면', style='OpH2')
    pw_change_desc = doc.add_paragraph()
    pw_change_desc.add_run('로그인된 사용자가 비밀번호를 변경하기 위한 화면입니다.\n\n')

    # 비밀번호 변경 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/change_password_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 6.3: 비밀번호 변경 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: change_password_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    pw_change_elements = doc.add_paragraph()
    pw_change_elements.add_run('6.3.1 화면 구성 요소\n').bold = True
    pw_change_elements.add_run('• 현재 비밀번호: 기존 비밀번호 확인용 입력 필드\n')
    pw_change_elements.add_run('• 새 비밀번호: 변경할 새 비밀번호 입력\n')
    pw_change_elements.add_run('• 새 비밀번호 확인: 새 비밀번호 재입력 검증\n')
    pw_change_elements.add_run('• 변경 버튼: 비밀번호 변경 실행\n\n')

    pw_change_usage = doc.add_paragraph()
    pw_change_usage.add_run('6.3.2 사용 방법\n').bold = True
    pw_change_usage.add_run('1. 현재 비밀번호 입력: 기존 비밀번호 확인\n')
    pw_change_usage.add_run('2. 새 비밀번호 입력: 새로운 비밀번호 입력\n')
    pw_change_usage.add_run('3. 새 비밀번호 확인: 동일한 새 비밀번호 재입력\n')
    pw_change_usage.add_run('4. 변경 버튼 클릭: 비밀번호 변경 완료\n\n')

    # 6.4 로그 관리
    doc.add_paragraph('7. 로그 관리', style='OpH1')

    # 6.1 로그 파일 위치 및 종류
    doc.add_paragraph('6.1 로그 파일 위치 및 종류', style='OpH2')
    log_files = doc.add_paragraph()
    log_files.add_run('6.1.1 로그 파일 종류\n').bold = True
    log_files.add_run('• 애플리케이션 로그 (app.log): 애플리케이션 실행 중 발생하는 이벤트 및 오류 기록\n')
    log_files.add_run('• 접근 로그 (access.log): 웹 서버 접근 기록\n')
    log_files.add_run('• 시스템 로그 (/var/log/syslog): 운영체제 수준의 시스템 이벤트\n\n')

    log_files.add_run('6.1.2 로그 파일 위치\n').bold = True
    log_files.add_run('• /home/msys/app/log/app.log\n')
    log_files.add_run('• /home/msys/app/log/access.log\n')
    log_files.add_run('• /var/log/syslog (시스템 로그)\n\n')

    # 6.2 로그 확인 방법
    doc.add_paragraph('6.2 로그 확인 방법', style='OpH2')
    log_check = doc.add_paragraph()
    log_check.add_run('6.2.1 실시간 로그 모니터링\n').bold = True
    log_check.add_run('로그 파일을 실시간으로 모니터링합니다:\n\n')
    log_check.add_run('tail -f /home/msys/app/log/app.log\n\n').font.name = 'Courier New'

    log_check.add_run('6.2.2 최근 로그 확인\n').bold = True
    log_check.add_run('최근 50줄의 로그를 확인합니다:\n\n')
    log_check.add_run('tail -50 /home/msys/app/log/app.log\n\n').font.name = 'Courier New'

    log_check.add_run('6.2.3 특정 시간대 로그 검색\n').bold = True
    log_check.add_run('특정 시간대의 로그를 검색합니다:\n\n')
    log_check.add_run('grep "2025-12-18" /home/msys/app/log/app.log\n\n').font.name = 'Courier New'

    log_check.add_run('6.2.4 오류 로그만 필터링\n').bold = True
    log_check.add_run('ERROR 레벨의 로그만 확인합니다:\n\n')
    log_check.add_run('grep "ERROR" /home/msys/app/log/app.log\n\n').font.name = 'Courier New'

    # 6.3 로그 데이터 해석
    doc.add_paragraph('6.3 로그 데이터 해석', style='OpH2')
    log_interpret = doc.add_paragraph()
    log_interpret.add_run('6.3.1 로그 레벨\n').bold = True
    log_interpret.add_run('• DEBUG: 상세한 디버깅 정보\n')
    log_interpret.add_run('• INFO: 일반적인 정보 메시지\n')
    log_interpret.add_run('• WARNING: 경고 상황\n')
    log_interpret.add_run('• ERROR: 오류 발생\n')
    log_interpret.add_run('• CRITICAL: 심각한 오류\n\n')

    log_interpret.add_run('6.3.2 일반적인 로그 메시지 예시\n').bold = True
    log_interpret.add_run('[2025-12-18 14:30:15] INFO: 애플리케이션 시작됨\n').font.name = 'Courier New'
    log_interpret.add_run('[2025-12-18 14:30:16] INFO: 데이터베이스 연결 성공\n').font.name = 'Courier New'
    log_interpret.add_run('[2025-12-18 14:35:22] ERROR: 데이터베이스 연결 실패 - Connection timeout\n').font.name = 'Courier New'
    log_interpret.add_run('[2025-12-18 14:35:23] WARNING: 메모리 사용률 85% 초과\n\n').font.name = 'Courier New'

    # 6.4 오류 발생 시 조치 방법
    doc.add_paragraph('6.4 오류 발생 시 조치 방법', style='OpH2')
    error_handling = doc.add_paragraph()
    error_handling.add_run('6.4.1 데이터베이스 연결 오류\n').bold = True
    error_handling.add_run('증상: "Connection timeout" 또는 "Connection refused" 오류\n')
    error_handling.add_run('조치:\n')
    error_handling.add_run('1. PostgreSQL 서비스 상태 확인: systemctl status postgresql\n')
    error_handling.add_run('2. 데이터베이스 서버 재시작: systemctl restart postgresql\n')
    error_handling.add_run('3. 네트워크 연결 확인: ping database_server\n')
    error_handling.add_run('4. .env 파일의 데이터베이스 설정 확인\n\n')

    error_handling.add_run('6.4.2 메모리 부족 오류\n').bold = True
    error_handling.add_run('증상: "MemoryError" 또는 메모리 사용률 90% 이상\n')
    error_handling.add_run('조치:\n')
    error_handling.add_run('1. 현재 메모리 사용량 확인: free -h\n')
    error_handling.add_run('2. 메모리 사용 프로세스 확인: ps aux --sort=-%mem | head\n')
    error_handling.add_run('3. 불필요한 프로세스 종료\n')
    error_handling.add_run('4. 서버 메모리 증설 고려\n\n')

    error_handling.add_run('6.4.3 디스크 공간 부족\n').bold = True
    error_handling.add_run('증상: "No space left on device" 오류\n')
    error_handling.add_run('조치:\n')
    error_handling.add_run('1. 디스크 사용량 확인: df -h\n')
    error_handling.add_run('2. 큰 파일 찾기: find / -type f -size +100M\n')
    error_handling.add_run('3. 로그 파일 정리: 로그 로테이션 실행\n')
    error_handling.add_run('4. 불필요한 파일 삭제\n\n')

    error_handling.add_run('6.4.4 애플리케이션 응답 없음\n').bold = True
    error_handling.add_run('증상: 웹 페이지가 로드되지 않음\n')
    error_handling.add_run('조치:\n')
    error_handling.add_run('1. 프로세스 상태 확인: ps aux | grep python\n')
    error_handling.add_run('2. 애플리케이션 로그 확인\n')
    error_handling.add_run('3. 애플리케이션 재시작\n')
    error_handling.add_run('4. 시스템 리소스 확인\n\n')

    # 6.5 로그 로테이션
    doc.add_paragraph('6.5 로그 로테이션', style='OpH2')
    log_rotate = doc.add_paragraph()
    log_rotate.add_run('6.5.1 로그 로테이션 설정\n').bold = True
    log_rotate.add_run('/etc/logrotate.d/msys 파일을 생성합니다:\n\n')
    log_rotate.add_run('/home/msys/app/log/*.log {\n').font.name = 'Courier New'
    log_rotate.add_run('    daily\n').font.name = 'Courier New'
    log_rotate.add_run('    rotate 30\n').font.name = 'Courier New'
    log_rotate.add_run('    compress\n').font.name = 'Courier New'
    log_rotate.add_run('    missingok\n').font.name = 'Courier New'
    log_rotate.add_run('    postrotate\n').font.name = 'Courier New'
    log_rotate.add_run('        systemctl reload msys\n').font.name = 'Courier New'
    log_rotate.add_run('    endscript\n').font.name = 'Courier New'
    log_rotate.add_run('}\n\n').font.name = 'Courier New'

    log_rotate.add_run('6.5.2 수동 로테이션 실행\n').bold = True
    log_rotate.add_run('logrotate -f /etc/logrotate.d/msys\n\n').font.name = 'Courier New'

    doc.save('MSYS_Operation_Manual_v2.docx')
    print('운영 매뉴얼 생성 완료: MSYS_Operation_Manual_v2.docx')

def create_database_manual():
    """DB 관리자용 데이터베이스 매뉴얼 생성"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('DBTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('DBH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('DBH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('DBH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    code_style = doc.styles.add_style('DBCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)

    # 표지
    title = doc.add_paragraph('MSYS 데이터베이스 매뉴얼', style='DBTitle')
    subtitle = doc.add_paragraph('DB 스키마 및 쿼리 참조 가이드', style='DBTitle')
    version = doc.add_paragraph('버전 1.14.2', style='DBTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='DBH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 데이터베이스 개요 ........................ 3\n')
    toc.add_run('2. 접속 정보 ............................. 4\n')
    toc.add_run('3. 테이블 스키마 ........................ 5\n')
    toc.add_run('4. 주요 쿼리 패턴 ........................ 15\n')
    toc.add_run('5. 인덱스 및 최적화 ...................... 20\n')
    toc.add_run('6. 백업 및 복원 ........................ 23\n')
    toc.add_run('부록 A. DDL 스크립트 ...................... 25\n')

    doc.add_page_break()

    # 1. 데이터베이스 개요
    doc.add_paragraph('1. 데이터베이스 개요', style='DBH1')
    overview = doc.add_paragraph()
    overview.add_run('1.1 시스템 개요\n').bold = True
    overview.add_run('MSYS는 데이터 수집 및 모니터링 시스템으로, PostgreSQL 데이터베이스를 사용하여 다양한 데이터를 저장하고 관리합니다.\n\n')

    overview.add_run('1.2 주요 기능\n').bold = True
    overview.add_run('• 데이터 수집 이력 관리\n')
    overview.add_run('• 사용자 및 권한 관리\n')
    overview.add_run('• 시스템 설정 관리\n')
    overview.add_run('• 분석 데이터 저장\n\n')

    overview.add_run('1.3 데이터베이스 버전\n').bold = True
    overview.add_run('• PostgreSQL 13 이상\n')
    overview.add_run('• 타임존: Asia/Seoul\n')
    overview.add_run('• 캐릭터셋: UTF-8\n\n')

    # 2. 접속 정보
    doc.add_paragraph('2. 접속 정보', style='DBH1')
    conn_info = doc.add_paragraph()
    conn_info.add_run('2.1 기본 접속 정보\n').bold = True
    conn_info.add_run('• 호스트: 10.200.153.136\n')
    conn_info.add_run('• 포트: 22543\n')
    conn_info.add_run('• 데이터베이스: etl_db_dev\n')
    conn_info.add_run('• 사용자: etl_user\n')
    conn_info.add_run('• 비밀번호: etl_password\n\n')

    conn_info.add_run('2.2 접속 예시\n').bold = True
    conn_info.add_run('psql -h 10.200.153.136 -p 22543 -U etl_user -d etl_db_dev\n\n').font.name = 'Courier New'

    # 3. 테이블 스키마
    doc.add_paragraph('3. 테이블 스키마', style='DBH1')

    # 3.1 tb_con_hist
    doc.add_paragraph('3.1 tb_con_hist - 데이터 수집 이력', style='DBH2')
    hist_desc = doc.add_paragraph()
    hist_desc.add_run('데이터 수집 작업의 상세 이력을 저장하는 메인 테이블입니다.\n\n')

    hist_desc.add_run('컬럼 상세:\n').bold = True
    hist_table = doc.add_table(rows=1, cols=6)
    hist_table.style = 'Table Grid'
    hdr_cells = hist_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    # 데이터 행 추가
    rows = [
        ['job_id', 'varchar', '250', 'NOT NULL', '', '작업 고유 식별자'],
        ['con_id', 'varchar', '250', 'NOT NULL', '', '연결 고유 식별자'],
        ['rqs_info', 'text', '', 'NULL', '', '요청 정보'],
        ['start_dt', 'timestamptz', '', 'NULL', '', '수집 시작 일시'],
        ['execution_dt', 'timestamptz', '', 'NULL', '', '실행 일시'],
        ['end_dt', 'timestamptz', '', 'NULL', '', '수집 종료 일시'],
        ['status', 'varchar', '20', 'NULL', '', '수집 상태'],
        ['trbl_hist_no', 'integer', '', 'NULL', '', '문제 이력 번호']
    ]

    for row_data in rows:
        row_cells = hist_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    hist_desc.add_run('\n\n상태 코드 값:\n').bold = True
    hist_desc.add_run('• CD901: 정상(성공) - 데이터 수집 성공\n')
    hist_desc.add_run('• CD902: 실패 - 데이터 수집 실패\n')
    hist_desc.add_run('• CD903: 미수집 - 예정된 수집 누락\n')
    hist_desc.add_run('• CD904: 측정중 - 수집 진행 중\n\n')

    hist_desc.add_run('제약조건:\n').bold = True
    hist_desc.add_run('• 프라이머리 키: (con_id, job_id)\n')
    hist_desc.add_run('• 트리거: trg_log_con_hist_changes (변경 로그 기록)\n\n')

    # 3.2 tb_user
    doc.add_paragraph('3.2 tb_user - 사용자 정보', style='DBH2')
    user_desc = doc.add_paragraph()
    user_desc.add_run('시스템 사용자의 계정 정보를 저장하는 테이블입니다.\n\n')

    user_table = doc.add_table(rows=1, cols=6)
    user_table.style = 'Table Grid'
    hdr_cells = user_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    user_rows = [
        ['user_id', 'varchar', '50', 'NOT NULL', '', '사용자 고유 ID'],
        ['user_pwd', 'varchar', '255', 'NOT NULL', '', '해시된 비밀번호'],
        ['acc_sts', 'varchar', '20', 'NOT NULL', "'PENDING'", '계정 상태'],
        ['acc_cre_dt', 'timestamptz', '', 'NULL', 'CURRENT_TIMESTAMP', '계정 생성 일시'],
        ['acc_apr_dt', 'timestamptz', '', 'NULL', '', '계정 승인 일시']
    ]

    for row_data in user_rows:
        row_cells = user_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    user_desc.add_run('\n\n계정 상태 코드:\n').bold = True
    user_desc.add_run('• PENDING: 승인 대기 - 회원가입 후 관리자 승인 필요\n')
    user_desc.add_run('• APPROVED: 승인 완료 - 정상 사용 가능\n')
    user_desc.add_run('• REJECTED: 승인 거부 - 사용 불가\n')
    user_desc.add_run('• SUSPENDED: 정지 - 일시적 사용 제한\n\n')

    # 3.3 tb_con_mst
    doc.add_paragraph('3.3 tb_con_mst - 마스터 코드', style='DBH2')
    mst_desc = doc.add_paragraph()
    mst_desc.add_run('시스템에서 사용하는 각종 코드와 설정값을 저장하는 마스터 테이블입니다.\n\n')

    mst_table = doc.add_table(rows=1, cols=6)
    mst_table.style = 'Table Grid'
    hdr_cells = mst_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    mst_rows = [
        ['cd_cl', 'varchar', '20', 'NOT NULL', '', '코드 분류'],
        ['cd', 'varchar', '20', 'NOT NULL', '', '코드 값'],
        ['cd_nm', 'varchar', '20', 'NULL', '', '코드명'],
        ['cd_desc', 'varchar', '50', 'NULL', '', '코드 설명'],
        ['item1', 'varchar', '50', 'NULL', '', '추가 항목 1'],
        ['item2', 'varchar', '50', 'NULL', '', '추가 항목 2'],
        ['item3', 'varchar', '150', 'NULL', '', '추가 항목 3'],
        ['item4', 'varchar', '50', 'NULL', '', '추가 항목 4'],
        ['item5', 'varchar', '50', 'NULL', '', '추가 항목 5'],
        ['item6', 'varchar', '50', 'NULL', '', '추가 항목 6'],
        ['item7', 'varchar', '50', 'NULL', '', '추가 항목 7'],
        ['item8', 'varchar', '50', 'NULL', '', '추가 항목 8'],
        ['item9', 'varchar', '400', 'NULL', '', '추가 항목 9'],
        ['item10', 'varchar', '400', 'NULL', '', '추가 항목 10'],
        ['update_dt', 'timestamptz', '', 'NULL', '', '수정 일시'],
        ['del_dt', 'timestamptz', '', 'NULL', '', '삭제 일시'],
        ['use_yn', 'char', '18', 'NULL', '', '사용 여부']
    ]

    for row_data in mst_rows:
        row_cells = mst_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    mst_desc.add_run('\n\n주요 코드 분류:\n').bold = True
    mst_desc.add_run('• CD900: 장애 코드 (item1: 영문명)\n')
    mst_desc.add_run('• CD100: 작업 유형 코드\n')
    mst_desc.add_run('• CD200: 상태 코드\n\n')

    # 3.4 tb_menu
    doc.add_paragraph('3.4 tb_menu - 메뉴 정보', style='DBH2')
    menu_desc = doc.add_paragraph()
    menu_desc.add_run('시스템의 메뉴 구조를 정의하는 테이블입니다.\n\n')

    menu_table = doc.add_table(rows=1, cols=6)
    menu_table.style = 'Table Grid'
    hdr_cells = menu_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    menu_rows = [
        ['menu_id', 'varchar', '50', 'NOT NULL', '', '메뉴 고유 ID'],
        ['menu_nm', 'varchar', '100', 'NOT NULL', '', '메뉴 표시 이름'],
        ['menu_url', 'varchar', '255', 'NOT NULL', '', '메뉴 URL 경로'],
        ['menu_order', 'integer', '', 'NULL', '', '메뉴 표시 순서']
    ]

    for row_data in menu_rows:
        row_cells = menu_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    menu_desc.add_run('\n\n제약조건:\n').bold = True
    menu_desc.add_run('• 프라이머리 키: menu_id\n\n')

    # 3.5 tb_mngr_sett
    doc.add_paragraph('3.5 tb_mngr_sett - 관리자 설정', style='DBH2')
    sett_desc = doc.add_paragraph()
    sett_desc.add_run('Job별 모니터링 임계값과 표시 설정을 저장하는 테이블입니다.\n\n')

    sett_table = doc.add_table(rows=1, cols=6)
    sett_table.style = 'Table Grid'
    hdr_cells = sett_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    sett_rows = [
        ['cd', 'varchar', '50', 'NOT NULL', '', '설정 대상 Job ID'],
        ['cnn_failr_thrs_val', 'integer', '', 'NULL', '3', '연속 실패 임계값'],
        ['cnn_warn_thrs_val', 'integer', '', 'NULL', '2', '연속 경고 임계값'],
        ['cnn_failr_icon_id', 'integer', '', 'NULL', '', '실패 상태 아이콘 ID'],
        ['cnn_failr_wrd_colr', 'varchar', '7', 'NULL', "'#FF0000'", '실패 상태 색상'],
        ['cnn_warn_icon_id', 'integer', '', 'NULL', '', '경고 상태 아이콘 ID'],
        ['cnn_warn_wrd_colr', 'varchar', '7', 'NULL', "'#FFA500'", '경고 상태 색상'],
        ['cnn_sucs_icon_id', 'integer', '', 'NULL', '', '성공 상태 아이콘 ID'],
        ['cnn_sucs_wrd_colr', 'varchar', '7', 'NULL', "'#008000'", '성공 상태 색상'],
        ['dly_sucs_rt_thrs_val', 'integer', '', 'NULL', '80', '일별 성공률 임계값(%)'],
        ['dd7_sucs_rt_thrs_val', 'integer', '', 'NULL', '75', '7일 성공률 임계값(%)'],
        ['mthl_sucs_rt_thrs_val', 'integer', '', 'NULL', '70', '월별 성공률 임계값(%)'],
        ['chrt_dsp_yn', 'boolean', '', 'NULL', 'true', '차트 표시 여부']
    ]

    for row_data in sett_rows:
        row_cells = sett_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    sett_desc.add_run('\n\n제약조건:\n').bold = True
    sett_desc.add_run('• 프라이머리 키: cd\n')
    sett_desc.add_run('• 외래키: cd → tb_con_mst.cd (Job ID 참조)\n\n')

    # 3.6 tb_data_spec
    doc.add_paragraph('3.6 tb_data_spec - 데이터 명세서', style='DBH2')
    spec_desc = doc.add_paragraph()
    spec_desc.add_run('외부 API의 명세 정보를 저장하는 테이블입니다.\n\n')

    spec_table = doc.add_table(rows=1, cols=6)
    spec_table.style = 'Table Grid'
    hdr_cells = spec_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    spec_rows = [
        ['id', 'integer', '', 'NOT NULL', 'nextval()', '고유 식별자'],
        ['data_name', 'varchar', '255', 'NOT NULL', '', '데이터 명칭'],
        ['description', 'text', '', 'NULL', '', '데이터 설명'],
        ['api_url', 'varchar', '2048', 'NULL', '', 'API 엔드포인트 URL'],
        ['provider', 'varchar', '255', 'NULL', '', '데이터 제공 기관'],
        ['keywords', 'varchar', '1024', 'NULL', '', '검색 키워드'],
        ['reference_doc_url', 'varchar', '2048', 'NULL', '', '참고 문서 URL'],
        ['created_at', 'timestamptz', '', 'NULL', 'CURRENT_TIMESTAMP', '생성 일시'],
        ['updated_at', 'timestamptz', '', 'NULL', 'CURRENT_TIMESTAMP', '수정 일시']
    ]

    for row_data in spec_rows:
        row_cells = spec_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    spec_desc.add_run('\n\n제약조건:\n').bold = True
    spec_desc.add_run('• 프라이머리 키: id\n')
    spec_desc.add_run('• 시퀀스: tb_data_spec_id_seq\n\n')

    # 3.7 tb_icon
    doc.add_paragraph('3.7 tb_icon - 아이콘 정보', style='DBH2')
    icon_desc = doc.add_paragraph()
    icon_desc.add_run('시스템에서 사용하는 아이콘 정보를 저장하는 테이블입니다.\n\n')

    icon_table = doc.add_table(rows=1, cols=6)
    icon_table.style = 'Table Grid'
    hdr_cells = icon_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    icon_rows = [
        ['icon_id', 'integer', '', 'NOT NULL', 'nextval()', '아이콘 고유 ID'],
        ['icon_cd', 'text', '', 'NOT NULL', '', '아이콘 코드'],
        ['icon_nm', 'varchar', '50', 'NOT NULL', '', '아이콘 이름'],
        ['icon_expl', 'varchar', '255', 'NULL', '', '아이콘 설명'],
        ['icon_cre_dt', 'timestamp', '', 'NULL', 'CURRENT_TIMESTAMP', '생성 일시'],
        ['icon_dsp_yn', 'boolean', '', 'NULL', 'true', '표시 여부']
    ]

    for row_data in icon_rows:
        row_cells = icon_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    icon_desc.add_run('\n\n제약조건:\n').bold = True
    icon_desc.add_run('• 프라이머리 키: icon_id\n')
    icon_desc.add_run('• 유니크: icon_cd\n')
    icon_desc.add_run('• 시퀀스: tb_icons_icon_id_seq\n\n')

    # 3.8 tb_user_auth_ctrl
    doc.add_paragraph('3.8 tb_user_auth_ctrl - 사용자 권한 제어', style='DBH2')
    auth_desc = doc.add_paragraph()
    auth_desc.add_run('사용자별 메뉴 접근 권한을 제어하는 테이블입니다.\n\n')

    auth_table = doc.add_table(rows=1, cols=6)
    auth_table.style = 'Table Grid'
    hdr_cells = auth_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    auth_rows = [
        ['auth_id', 'integer', '', 'NOT NULL', 'nextval()', '권한 고유 ID'],
        ['user_id', 'varchar', '50', 'NOT NULL', '', '사용자 ID'],
        ['menu_id', 'varchar', '50', 'NOT NULL', '', '메뉴 ID'],
        ['auth_yn', 'boolean', '', 'NOT NULL', 'true', '권한 여부']
    ]

    for row_data in auth_rows:
        row_cells = auth_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    auth_desc.add_run('\n\n제약조건:\n').bold = True
    auth_desc.add_run('• 프라이머리 키: auth_id\n')
    auth_desc.add_run('• 유니크: (user_id, menu_id)\n')
    auth_desc.add_run('• 외래키: user_id → tb_user.user_id\n')
    auth_desc.add_run('• 외래키: menu_id → tb_menu.menu_id\n')
    auth_desc.add_run('• 시퀀스: tb_user_access_control_id_seq\n\n')

    # 3.9 tb_user_data_perm_auth_ctrl
    doc.add_paragraph('3.9 tb_user_data_perm_auth_ctrl - 데이터 접근 권한', style='DBH2')
    data_perm_desc = doc.add_paragraph()
    data_perm_desc.add_run('사용자별 Job ID 데이터 접근 권한을 제어하는 테이블입니다.\n\n')

    data_perm_table = doc.add_table(rows=1, cols=6)
    data_perm_table.style = 'Table Grid'
    hdr_cells = data_perm_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    data_perm_rows = [
        ['perm_id', 'integer', '', 'NOT NULL', 'nextval()', '권한 고유 ID'],
        ['user_id', 'varchar', '50', 'NOT NULL', '', '사용자 ID'],
        ['job_id', 'varchar', '50', 'NOT NULL', '', 'Job ID'],
        ['perm_yn', 'boolean', '', 'NOT NULL', 'true', '접근 권한 여부']
    ]

    for row_data in data_perm_rows:
        row_cells = data_perm_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    data_perm_desc.add_run('\n\n제약조건:\n').bold = True
    data_perm_desc.add_run('• 프라이머리 키: perm_id\n')
    data_perm_desc.add_run('• 유니크: (user_id, job_id)\n')
    data_perm_desc.add_run('• 외래키: user_id → tb_user.user_id\n')
    data_perm_desc.add_run('• 외래키: job_id → tb_con_mst.cd\n\n')

    # 3.10 기타 주요 테이블 개요
    doc.add_paragraph('3.10 기타 주요 테이블 개요', style='DBH2')
    other_tables = doc.add_paragraph()
    other_tables.add_run('시스템에서 사용하는 기타 주요 테이블들:\n\n').bold = True

    other_tables.add_run('• tb_con_hist_evnt_log: 수집 이력 변경 로그\n')
    other_tables.add_run('• tb_con_trbl_hist: 문제 발생 이력\n')
    other_tables.add_run('• tb_data_clt_schd_sett: 수집 스케줄 설정\n')
    other_tables.add_run('• tb_data_spec_parm: 데이터 명세서 파라미터\n')
    other_tables.add_run('• tb_user_acs_log: 사용자 접근 로그\n')
    other_tables.add_run('• tb_col_mapp: 컬럼 매핑 정보\n\n')

    # 4. 주요 쿼리 패턴
    doc.add_paragraph('4. 주요 쿼리 패턴', style='DBH1')

    # 4.1 성공률 계산 쿼리
    doc.add_paragraph('4.1 성공률 계산 쿼리', style='DBH2')
    success_query = doc.add_paragraph()
    success_query.add_run('일별/Job별 성공률을 계산하는 표준 쿼리 패턴:\n\n')
    success_query.add_run('SELECT\n')
    success_query.add_run('    (start_dt AT TIME ZONE \'Asia/Seoul\')::date AS date,\n')
    success_query.add_run('    job_id,\n')
    success_query.add_run('    (SUM(CASE WHEN status = \'CD901\' THEN 1 ELSE 0 END) * 100.0 / NULLIF(COUNT(*), 0)) AS success_rate\n')
    success_query.add_run('FROM tb_con_hist\n')
    success_query.add_run('{where_clause}\n')
    success_query.add_run('GROUP BY (start_dt AT TIME ZONE \'Asia/Seoul\')::date, job_id\n')
    success_query.add_run('ORDER BY (start_dt AT TIME ZONE \'Asia/Seoul\')::date, job_id;\n\n').font.name = 'Courier New'

    success_query.add_run('계산 방식:\n').bold = True
    success_query.add_run('• 성공률 = (성공 건수 ÷ 전체 건수) × 100\n')
    success_query.add_run('• 성공 기준: status = \'CD901\'\n')
    success_query.add_run('• 시간대: Asia/Seoul 기준\n\n')

    # 4.2 집계 쿼리 패턴
    doc.add_paragraph('4.2 집계 쿼리 패턴', style='DBH2')
    agg_query = doc.add_paragraph()
    agg_query.add_run('다양한 기간별 통계를 계산하는 집계 쿼리:\n\n')
    agg_query.add_run('-- 일별 집계\n')
    agg_query.add_run('SELECT DATE_TRUNC(\'day\', start_dt AT TIME ZONE \'Asia/Seoul\') AS period,\n')
    agg_query.add_run('       COUNT(*) as total_count,\n')
    agg_query.add_run('       SUM(CASE WHEN status = \'CD901\' THEN 1 ELSE 0 END) as success_count\n')
    agg_query.add_run('FROM tb_con_hist\n')
    agg_query.add_run('WHERE start_dt >= \'2025-01-01\'\n')
    agg_query.add_run('GROUP BY DATE_TRUNC(\'day\', start_dt AT TIME ZONE \'Asia/Seoul\')\n')
    agg_query.add_run('ORDER BY period;\n\n').font.name = 'Courier New'

    # 5. 인덱스 및 최적화
    doc.add_paragraph('5. 인덱스 및 최적화', style='DBH1')
    index_desc = doc.add_paragraph()
    index_desc.add_run('5.1 주요 인덱스\n').bold = True
    index_desc.add_run('• tb_con_hist: (con_id, job_id) - 프라이머리 키\n')
    index_desc.add_run('• tb_con_hist: start_dt - 시간 범위 쿼리 최적화\n')
    index_desc.add_run('• tb_con_hist: (job_id, start_dt) - Job별 기간 쿼리\n')
    index_desc.add_run('• tb_user: user_id - 프라이머리 키\n')
    index_desc.add_run('• tb_con_mst: (cd_cl, cd) - 프라이머리 키\n\n')

    index_desc.add_run('5.2 쿼리 최적화 팁\n').bold = True
    index_desc.add_run('• 시간 범위 쿼리: start_dt 인덱스 활용\n')
    index_desc.add_run('• Job 필터링: job_id + start_dt 복합 인덱스\n')
    index_desc.add_run('• 대량 데이터: LIMIT과 OFFSET 적절히 사용\n')
    index_desc.add_run('• 파티셔닝: 대용량 테이블은 날짜별 파티셔닝 고려\n\n')

    # 6. 백업 및 복원
    doc.add_paragraph('6. 백업 및 복원', style='DBH1')
    backup_desc = doc.add_paragraph()
    backup_desc.add_run('6.1 전체 백업\n').bold = True
    backup_desc.add_run('pg_dump -h 10.200.153.136 -p 22543 -U etl_user -d etl_db_dev > backup.sql\n\n').font.name = 'Courier New'

    backup_desc.add_run('6.2 복원\n').bold = True
    backup_desc.add_run('psql -h 10.200.153.136 -p 22543 -U etl_user -d etl_db_dev < backup.sql\n\n').font.name = 'Courier New'

    backup_desc.add_run('6.3 자동 백업 스크립트\n').bold = True
    backup_desc.add_run('# crontab에 추가\n')
    backup_desc.add_run('0 2 * * * pg_dump -U etl_user etl_db_dev > /backup/msys_$(date +\\%Y\\%m\\%d).sql\n\n').font.name = 'Courier New'

    doc.save('MSYS_Database_Manual.docx')
    print('데이터베이스 매뉴얼 생성 완료: MSYS_Database_Manual.docx')

def create_cpu_monitor_manual():
    """CPU 모니터링 앱 설명서 생성"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('CpuTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('CpuH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('CpuH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('CpuH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    code_style = doc.styles.add_style('CpuCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)

    # 표지
    title = doc.add_paragraph('MSYS CPU 모니터링 앱 설명서', style='CpuTitle')
    subtitle = doc.add_paragraph('시스템 리소스 모니터링 가이드', style='CpuTitle')
    version = doc.add_paragraph('버전 1.0', style='CpuTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='CpuH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 개요 ........................ 3\n')
    toc.add_run('2. 설치 및 실행 ................ 4\n')
    toc.add_run('3. 모니터링 기능 ................ 6\n')
    toc.add_run('4. 설정 옵션 .................. 9\n')
    toc.add_run('5. 로그 및 결과 해석 ........... 11\n')
    toc.add_run('6. 문제 해결 .................. 14\n')
    toc.add_run('부록 A. 기술 사양 ................ 16\n')

    doc.add_page_break()

    # 1. 개요
    doc.add_paragraph('1. 개요', style='CpuH1')
    overview = doc.add_paragraph()
    overview.add_run('1.1 목적\n').bold = True
    overview.add_run('MSYS CPU 모니터링 앱은 시스템의 CPU 사용률을 실시간으로 모니터링하고, 임계값 초과 시 경고를 기록하는 도구입니다. 시스템 성능 저하의 조기 발견과 문제 해결을 지원합니다.\n\n')

    overview.add_run('1.2 주요 기능\n').bold = True
    overview.add_run('• 실시간 CPU 사용률 모니터링\n')
    overview.add_run('• 임계값 기반 경고 시스템\n')
    overview.add_run('• 급격한 CPU 사용량 증가 감지\n')
    overview.add_run('• 상위 CPU 사용 프로세스 식별\n')
    overview.add_run('• 자동 로그 기록 및 로테이션\n\n')

    overview.add_run('1.3 지원 플랫폼\n').bold = True
    overview.add_run('• Windows (pythonw.exe 사용)\n')
    overview.add_run('• Linux/macOS (nohup 사용)\n')
    overview.add_run('• 백그라운드 실행 지원\n\n')

    # 2. 설치 및 실행
    doc.add_paragraph('2. 설치 및 실행', style='CpuH1')

    # 2.1 요구사항
    doc.add_paragraph('2.1 시스템 요구사항', style='CpuH2')
    reqs = doc.add_paragraph()
    reqs.add_run('• Python 3.8 이상\n')
    reqs.add_run('• psutil 라이브러리 (pip install psutil)\n')
    reqs.add_run('• 로그 파일 쓰기 권한\n\n')

    # 2.2 설치
    doc.add_paragraph('2.2 설치', style='CpuH2')
    install = doc.add_paragraph()
    install.add_run('필수 라이브러리 설치:\n\n')
    install.add_run('pip install psutil\n\n').font.name = 'Courier New'

    # 2.3 실행
    doc.add_paragraph('2.3 실행 방법', style='CpuH2')
    run = doc.add_paragraph()
    run.add_run('2.3.1 백그라운드 실행 (권장)\n').bold = True
    run.add_run('python start_cpu_monitor.py\n\n').font.name = 'Courier New'

    run.add_run('2.3.2 직접 실행\n').bold = True
    run.add_run('python utils/cpu_monitor.py\n\n').font.name = 'Courier New'

    run.add_run('2.3.3 테스트 모드 실행\n').bold = True
    run.add_run('python utils/cpu_monitor.py test\n\n').font.name = 'Courier New'

    # 2.4 실행 확인
    doc.add_paragraph('2.4 실행 상태 확인', style='CpuH2')
    check = doc.add_paragraph()
    check.add_run('프로세스 확인 (Windows):\n\n')
    check.add_run('tasklist /FI "IMAGENAME eq python.exe"\n\n').font.name = 'Courier New'

    check.add_run('프로세스 확인 (Linux):\n\n')
    check.add_run('ps aux | grep python\n\n').font.name = 'Courier New'

    # 3. 모니터링 기능
    doc.add_paragraph('3. 모니터링 기능', style='CpuH1')

    # 3.1 기본 모니터링
    doc.add_paragraph('3.1 기본 CPU 모니터링', style='CpuH2')
    basic = doc.add_paragraph()
    basic.add_run('• 5초 간격으로 CPU 사용률 측정\n')
    basic.add_run('• 실시간 로그 기록\n')
    basic.add_run('• 백분율 단위 표시\n\n')

    # 3.2 임계값 모니터링
    doc.add_paragraph('3.2 임계값 기반 경고', style='CpuH2')
    threshold = doc.add_paragraph()
    threshold.add_run('기본 임계값:\n').bold = True
    threshold.add_run('• CPU 사용률 80% 초과 시 경고\n')
    threshold.add_run('• 이전 값 대비 20% 급격한 증가 시 경고\n\n')

    # 3.3 프로세스 모니터링
    doc.add_paragraph('3.3 상위 CPU 사용 프로세스', style='CpuH2')
    process = doc.add_paragraph()
    process.add_run('CPU 사용률이 임계값을 초과하면 상위 5개 프로세스 정보를 기록:\n').bold = True
    process.add_run('• 프로세스 ID (PID)\n')
    process.add_run('• 프로세스 이름\n')
    process.add_run('• CPU 사용률\n\n')

    # 3.4 급격한 변화 감지
    doc.add_paragraph('3.4 급격한 CPU 변화 감지', style='CpuH2')
    spike = doc.add_paragraph()
    spike.add_run('이전 측정값과 비교하여 급격한 증가를 감지:\n').bold = True
    spike.add_run('• 기본 임계값: 20% 증가\n')
    spike.add_run('• 스팸 방지: 10초 쿨다운\n\n')

    # 4. 설정 옵션
    doc.add_paragraph('4. 설정 옵션', style='CpuH1')

    # 4.1 모니터링 간격
    doc.add_paragraph('4.1 모니터링 간격 설정', style='CpuH2')
    interval = doc.add_paragraph()
    interval.add_run('interval 매개변수로 설정 (기본값: 5초)\n\n')
    interval.add_run('monitor_cpu(interval=10)  # 10초 간격\n\n').font.name = 'Courier New'

    # 4.2 임계값 설정
    doc.add_paragraph('4.2 임계값 설정', style='CpuH2')
    thresholds = doc.add_paragraph()
    thresholds.add_run('threshold: CPU 사용률 임계값 (기본값: 80%)\n')
    thresholds.add_run('spike_threshold: 급격한 증가 임계값 (기본값: 20%)\n\n')
    thresholds.add_run('monitor_cpu(threshold=90, spike_threshold=15)\n\n').font.name = 'Courier New'

    # 4.3 테스트 모드
    doc.add_paragraph('4.3 테스트 모드', style='CpuH2')
    test_mode = doc.add_paragraph()
    test_mode.add_run('test_mode=True로 설정 시 3번만 실행하고 종료\n\n')
    test_mode.add_run('monitor_cpu(test_mode=True)\n\n').font.name = 'Courier New'

    # 5. 로그 및 결과 해석
    doc.add_paragraph('5. 로그 및 결과 해석', style='CpuH1')

    # 5.1 로그 파일 위치
    doc.add_paragraph('5.1 로그 파일', style='CpuH2')
    log_location = doc.add_paragraph()
    log_location.add_run('• 기본 위치: log/cpu_monitor.log\n')
    log_location.add_run('• 로그 로테이션: 일별 자동 로테이션\n')
    log_location.add_run('• 압축: 이전 로그 자동 압축\n\n')

    # 5.2 로그 형식
    doc.add_paragraph('5.2 로그 형식', style='CpuH2')
    log_format = doc.add_paragraph()
    log_format.add_run('일반 로그:\n').bold = True
    log_format.add_run('2025-12-23 13:30:15 - INFO - 현재 CPU 사용률: 45.2%\n\n').font.name = 'Courier New'

    log_format.add_run('경고 로그:\n').bold = True
    log_format.add_run('2025-12-23 13:30:20 - WARNING - CPU 사용률 경고: 85.7%\n\n').font.name = 'Courier New'

    log_format.add_run('급격한 증가 로그:\n').bold = True
    log_format.add_run('2025-12-23 13:30:25 - WARNING - CPU 급격한 상승 감지: 45.2% → 85.7% (+40.5%)\n\n').font.name = 'Courier New'

    # 5.3 프로세스 정보 로그
    doc.add_paragraph('5.3 프로세스 정보 로그', style='CpuH2')
    proc_log = doc.add_paragraph()
    proc_log.add_run('상위 CPU 사용 프로세스:\n').bold = True
    proc_log.add_run('2025-12-23 13:30:30 - WARNING -   PID: 1234, 이름: python.exe, CPU: 45.2%\n')
    proc_log.add_run('2025-12-23 13:30:30 - WARNING -   PID: 5678, 이름: chrome.exe, CPU: 23.1%\n\n').font.name = 'Courier New'

    # 5.4 로그 분석
    doc.add_paragraph('5.4 로그 분석 방법', style='CpuH2')
    analysis = doc.add_paragraph()
    analysis.add_run('5.4.1 정상 패턴\n').bold = True
    analysis.add_run('• CPU 사용률 1-50%: 정상 작동\n')
    analysis.add_run('• 일정한 패턴 유지\n\n')

    analysis.add_run('5.4.2 경고 패턴\n').bold = True
    analysis.add_run('• CPU 사용률 80% 초과 지속\n')
    analysis.add_run('• 급격한 사용량 증가\n')
    analysis.add_run('• 상위 프로세스 확인 필요\n\n')

    analysis.add_run('5.4.3 문제 패턴\n').bold = True
    analysis.add_run('• CPU 사용률 95% 이상 지속\n')
    analysis.add_run('• 동일 프로세스의 반복적 고사용량\n')
    analysis.add_run('• 시스템 응답성 저하\n\n')

    # 6. 문제 해결
    doc.add_paragraph('6. 문제 해결', style='CpuH1')

    # 6.1 일반적인 문제
    doc.add_paragraph('6.1 일반적인 문제', style='CpuH2')
    issues = doc.add_paragraph()
    issues.add_run('6.1.1 모니터링이 시작되지 않음\n').bold = True
    issues.add_run('• psutil 라이브러리 설치 확인: pip install psutil\n')
    issues.add_run('• Python 경로 확인\n')
    issues.add_run('• 로그 디렉토리 쓰기 권한 확인\n\n')

    issues.add_run('6.1.2 로그 파일이 생성되지 않음\n').bold = True
    issues.add_run('• log 디렉토리 존재 확인\n')
    issues.add_run('• 파일 시스템 권한 확인\n')
    issues.add_run('• 디스크 공간 부족 확인\n\n')

    issues.add_run('6.1.3 CPU 사용률이 0%로 표시됨\n').bold = True
    issues.add_run('• 시스템 권한 확인 (관리자 권한 필요할 수 있음)\n')
    issues.add_run('• psutil 버전 호환성 확인\n\n')

    # 6.2 고급 문제 해결
    doc.add_paragraph('6.2 고급 문제 해결', style='CpuH2')
    advanced = doc.add_paragraph()
    advanced.add_run('6.2.1 백그라운드 프로세스 종료\n').bold = True
    advanced.add_run('Windows:\n')
    advanced.add_run('taskkill /F /IM python.exe\n\n').font.name = 'Courier New'
    advanced.add_run('Linux:\n')
    advanced.add_run('pkill -f cpu_monitor.py\n\n').font.name = 'Courier New'

    advanced.add_run('6.2.2 로그 분석 스크립트\n').bold = True
    advanced.add_run('grep "WARNING" log/cpu_monitor.log | tail -10\n\n').font.name = 'Courier New'

    # 부록 A. 기술 사양
    doc.add_paragraph('부록 A. 기술 사양', style='CpuH1')
    spec = doc.add_paragraph()
    spec.add_run('A.1 의존성 라이브러리\n').bold = True
    spec.add_run('• psutil >= 5.0.0: 시스템 정보 수집\n')
    spec.add_run('• Python >= 3.8: 런타임 환경\n\n')

    spec.add_run('A.2 시스템 요구사항\n').bold = True
    spec.add_run('• 메모리: 최소 50MB\n')
    spec.add_run('• 저장소: 최소 10MB (로그 포함)\n')
    spec.add_run('• CPU: 모니터링 오버헤드 최소\n\n')

    spec.add_run('A.3 제한사항\n').bold = True
    spec.add_run('• Windows: 관리자 권한 필요 (프로세스 정보 수집 시)\n')
    spec.add_run('• Linux: /proc 파일시스템 접근 권한 필요\n')
    spec.add_run('• macOS: 시스템 통합 권한 필요할 수 있음\n\n')

    doc.save('MSYS_CPU_Monitor_Manual.docx')
    print('CPU 모니터링 앱 설명서 생성 완료: MSYS_CPU_Monitor_Manual.docx')

if __name__ == '__main__':
    create_installation_manual()
    create_function_manual()
    create_operation_manual()
    create_database_manual()
    create_cpu_monitor_manual()
    print('모든 메뉴얼 생성이 완료되었습니다.')
